"""Utilidades python-docx para documentos encuadrados dentro de margenes.

Importacion tipica:
    from doc_helpers import (
        PortraitDoc, LandscapeDoc,
        add_table_safe, add_heading_box, add_body, add_callout, add_cita,
        USABLE_W_PORTRAIT, USABLE_W_LANDSCAPE, distribute_widths,
        NAVY, AMBAR, RED, GREEN, GREY,
        LIGHT_BG, NAVY_BG, AMBAR_BG, RED_BG, GREEN_BG,
    )

Doctrina:
- Toda tabla con `autofit = False` y anchos explicitos.
- Toda tabla con `alignment = WD_TABLE_ALIGNMENT.CENTER`.
- Anchos de columna suman al ancho util de la pagina (no mas).
- Las tablas mono-columna tambien con ancho explicito.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# === Paleta institucional ===
NAVY = RGBColor(0x1F, 0x2A, 0x44)
AMBAR = RGBColor(0xD9, 0x82, 0x1B)
RED = RGBColor(0x9C, 0x1A, 0x1A)
GREEN = RGBColor(0x0B, 0x70, 0x3C)
GREY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

LIGHT_BG = "F4F6FA"
NAVY_BG = "1F2A44"
AMBAR_BG = "FCE8B3"
RED_BG = "F5C9C2"
GREEN_BG = "C6F3DE"


# === Constantes de pagina A4 ===
PAGE_W_PORTRAIT = 21.0
PAGE_H_PORTRAIT = 29.7
PAGE_W_LANDSCAPE = 29.7
PAGE_H_LANDSCAPE = 21.0

# Margenes default (cm)
MARGIN_TOP = 2.0
MARGIN_BOTTOM = 2.0
MARGIN_LEFT = 2.0
MARGIN_RIGHT = 2.0

USABLE_W_PORTRAIT = PAGE_W_PORTRAIT - MARGIN_LEFT - MARGIN_RIGHT       # 17.0 cm
USABLE_H_PORTRAIT = PAGE_H_PORTRAIT - MARGIN_TOP - MARGIN_BOTTOM       # 25.7 cm
USABLE_W_LANDSCAPE = PAGE_W_LANDSCAPE - MARGIN_LEFT - MARGIN_RIGHT     # 25.7 cm
USABLE_H_LANDSCAPE = PAGE_H_LANDSCAPE - MARGIN_TOP - MARGIN_BOTTOM     # 17.0 cm


def _apply_margins(section, top=MARGIN_TOP, bottom=MARGIN_BOTTOM,
                   left=MARGIN_LEFT, right=MARGIN_RIGHT):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def PortraitDoc(margins=(MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT)):
    """Crea un Document A4 vertical con margenes encuadrados. Retorna (doc, usable_width_cm)."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(PAGE_W_PORTRAIT)
        section.page_height = Cm(PAGE_H_PORTRAIT)
        _apply_margins(section, *margins)
    usable_w = PAGE_W_PORTRAIT - margins[2] - margins[3]
    return doc, usable_w


def LandscapeDoc(margins=(MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT)):
    """Crea un Document A4 horizontal con margenes encuadrados. Retorna (doc, usable_width_cm)."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(PAGE_W_LANDSCAPE)
        section.page_height = Cm(PAGE_H_LANDSCAPE)
        _apply_margins(section, *margins)
    usable_w = PAGE_W_LANDSCAPE - margins[2] - margins[3]
    return doc, usable_w


def distribute_widths(usable_width_cm, ratios):
    """Distribuye ratios proporcionalmente sumando al ancho util exacto.

    Ejemplo: distribute_widths(17.0, [3, 2, 4]) → [5.667, 3.778, 7.556]
    """
    total = sum(ratios)
    return [usable_width_cm * r / total for r in ratios]


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, italic=False,
                  align=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text if text is not None else "")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if white:
        r.font.color.rgb = WHITE
    elif color is not None:
        r.font.color.rgb = color


def _remove_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'nil')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _font_size_by_cols(n_cols):
    if n_cols <= 4: return 10
    if n_cols <= 6: return 9
    if n_cols <= 10: return 8
    return 7


def add_table_safe(doc, rows, col_widths_cm, header_bg=NAVY_BG,
                   header_white=True, font_size=None, align=WD_TABLE_ALIGNMENT.CENTER):
    """Anade una tabla con anchos explicitos, centrada y con cabecera estilizada.

    rows: lista de listas. La PRIMERA fila se trata como encabezado y se pinta con header_bg.
    col_widths_cm: lista de anchos en cm. DEBE sumar al ancho util de la pagina.
    """
    if not rows:
        return None
    n_rows = len(rows)
    n_cols = len(rows[0])
    assert all(len(r) == n_cols for r in rows), "Todas las filas deben tener el mismo numero de columnas"
    assert len(col_widths_cm) == n_cols, "col_widths_cm debe coincidir con el numero de columnas"

    if font_size is None:
        font_size = _font_size_by_cols(n_cols)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False
    table.alignment = align

    # Anchos explicitos
    for i, w in enumerate(col_widths_cm):
        table.columns[i].width = Cm(w)
        for cell in table.columns[i].cells:
            cell.width = Cm(w)

    # Encabezado
    for j, h in enumerate(rows[0]):
        set_cell_bg(table.rows[0].cells[j], header_bg)
        set_cell_text(table.rows[0].cells[j], str(h), bold=True, white=header_white,
                      size=font_size)

    # Cuerpo
    for i, row_data in enumerate(rows[1:], start=1):
        for j, val in enumerate(row_data):
            set_cell_text(table.rows[i].cells[j], str(val) if val is not None else "",
                          size=font_size)
    return table


def add_single_col_block(doc, text, usable_w_cm, bg=None, bold=False, color=None,
                         size=10, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before_pt=4, space_after_pt=8):
    """Bloque de una sola celda al ancho util — util para callouts, citas, espacios para anotar."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(usable_w_cm)
    t.rows[0].cells[0].width = Cm(usable_w_cm)
    cell = t.rows[0].cells[0]
    if bg:
        set_cell_bg(cell, bg)
    set_cell_text(cell, text, bold=bold, color=color, size=size, italic=italic, align=align)
    cell.paragraphs[0].paragraph_format.space_before = Pt(space_before_pt)
    cell.paragraphs[0].paragraph_format.space_after = Pt(space_after_pt)
    return t


def add_heading_box(doc, text, usable_w_cm, color_bg=NAVY_BG, size=14):
    """Encabezado de seccion en banda de color, ancho util completo."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(usable_w_cm)
    t.rows[0].cells[0].width = Cm(usable_w_cm)
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, color_bg)
    set_cell_text(cell, text, bold=True, white=True, size=size,
                  align=WD_ALIGN_PARAGRAPH.LEFT)
    cell.paragraphs[0].paragraph_format.space_before = Pt(8)
    cell.paragraphs[0].paragraph_format.space_after = Pt(8)
    _remove_borders(cell)
    doc.add_paragraph()
    return t


def add_callout(doc, text, usable_w_cm, color_bg=AMBAR_BG, color_text=NAVY):
    """Callout/aviso destacado en color, ancho util completo."""
    return add_single_col_block(doc, text, usable_w_cm, bg=color_bg, color=color_text,
                                size=10, space_before_pt=4, space_after_pt=8)


def add_body(doc, text, size=10, italic=False, color=None, bold=False,
             align=WD_ALIGN_PARAGRAPH.LEFT):
    """Parrafo normal."""
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_cita(doc, text, usable_w_cm):
    """Linea de cita en pequena, gris e italica — anclada al ancho util para no desbordar."""
    return add_single_col_block(doc, "Fuente: " + text, usable_w_cm,
                                bg=None, color=GREY, size=8, italic=True,
                                space_before_pt=2, space_after_pt=8)


def add_response_box(doc, label, usable_w_cm, n_blank_lines=4):
    """Caja para anotar respuesta en una reunion: cabecera + area en blanco, mono-columna."""
    # Cabecera
    t1 = doc.add_table(rows=1, cols=1)
    t1.autofit = False
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.columns[0].width = Cm(usable_w_cm)
    t1.rows[0].cells[0].width = Cm(usable_w_cm)
    cell = t1.rows[0].cells[0]
    set_cell_bg(cell, LIGHT_BG)
    set_cell_text(cell, label, bold=True, color=NAVY, size=9)

    # Area blanca
    t2 = doc.add_table(rows=1, cols=1)
    t2.autofit = False
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.columns[0].width = Cm(usable_w_cm)
    t2.rows[0].cells[0].width = Cm(usable_w_cm)
    set_cell_text(t2.rows[0].cells[0], "\n" * n_blank_lines, size=10)
    return (t1, t2)


def add_signature_row(doc, labels, usable_w_cm):
    """Fila de firmas distribuida en N columnas iguales al ancho util."""
    n = len(labels)
    col_w = usable_w_cm / n
    t = doc.add_table(rows=4, cols=n)
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(n):
        t.columns[i].width = Cm(col_w)
        for r in range(4):
            t.rows[r].cells[i].width = Cm(col_w)
    for i, l in enumerate(labels):
        set_cell_text(t.rows[0].cells[i], l, bold=True, color=NAVY, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(t.rows[1].cells[i], "\n\n\n", size=10)
        set_cell_text(t.rows[2].cells[i], "_______________________________",
                      size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(t.rows[3].cells[i], "Fecha: ____ / ____ / ______",
                      size=8, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    return t
