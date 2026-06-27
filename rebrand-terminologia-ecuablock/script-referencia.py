# -*- coding: utf-8 -*-
"""
REELABORACION INTEGRAL v2 - ETAPAS II y III
============================================
- Integra las enmiendas de la auditoria Py 7572/25
- Rebranding: EcuaBlock -> EcuaLedger Soberana
- NO toca LOFPD ni su Reglamento (Etapa I la hace el usuario manualmente)
- NO genera MD. Solo DOCX + PDF.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert as d2p_convert

BASE = Path(r"C:\Users\datos\Dropbox\var 91\Neg Inm\1 Proy Belen 2026"
            r"\Tokenizacion Activos Reales - RWA\JRPFM"
            r"\Proyecto Marco Regulatorio Asamblea")

PURPURA_ADN = RGBColor(0x72, 0x2F, 0xA1)

# ===================================================================
# REBRANDING - reemplazos terminologicos
# ===================================================================

REBRAND_RULES = [
    # Más específico primero
    (r"\bEcuaBlock\s*-\s*IBPP\b", "EcuaLedger Soberana - IBPP"),
    (r"\bEcuaBlock\b", "EcuaLedger Soberana"),
    (r"\becuablock\b", "EcuaLedger Soberana"),
    (r"\bEcuablock\b", "EcuaLedger Soberana"),
    (r"\bECUABLOCK\b", "ECUALEDGER SOBERANA"),
]

def rebrand(text: str) -> str:
    """Aplica el rebrand EcuaBlock -> EcuaLedger Soberana respetando word boundaries."""
    out = text
    for pat, repl in REBRAND_RULES:
        out = re.sub(pat, repl, out)
    return out


# ===================================================================
# UTILIDADES COMUNES DE FORMATO
# ===================================================================

def setup_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Cambria"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


def add_titulo(doc, txt, size=16, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = PURPURA_ADN


def add_h2(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = PURPURA_ADN


def add_h3(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(11.5)


def add_p(doc, txt, justify=True, italic=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(txt)
    if italic:
        r.italic = True


def add_articulo(doc, num_y_epigrafe, cuerpo):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(num_y_epigrafe + " ")
    r1.bold = True
    p.add_run(cuerpo)


def render_block_text(doc, text):
    for line in text.strip().split("\n"):
        ln = line.rstrip()
        if not ln.strip():
            doc.add_paragraph()
            continue
        if ln.startswith("# "):
            add_titulo(doc, ln[2:].strip(), size=15)
        elif ln.startswith("## "):
            add_h2(doc, ln[3:].strip())
        elif ln.startswith("### "):
            add_h3(doc, ln[4:].strip())
        elif re.match(r"^Artículo\s+[\dA-Z]+", ln):
            m = re.match(r"^(Artículo\s+[^.]+\.-)\s*(.*)$", ln)
            if m:
                add_articulo(doc, m.group(1), m.group(2))
            else:
                add_p(doc, ln)
        elif re.match(r"^(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA|DÉCIMA|UNICA|ÚNICA)\.?\-", ln):
            m = re.match(r"^([^.]+\.-)\s*(.*)$", ln)
            if m:
                add_articulo(doc, m.group(1), m.group(2))
            else:
                add_p(doc, ln)
        elif ln.startswith("- ") or ln.startswith("* "):
            doc.add_paragraph(ln[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", ln):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", ln), style="List Number")
        elif re.match(r"^[a-z]\)\s", ln):
            add_p(doc, ln)
        else:
            add_p(doc, ln)


def save_docx_pdf(doc, base_name):
    docx_path = BASE / f"{base_name}.docx"
    pdf_path = BASE / f"{base_name}.pdf"
    doc.save(docx_path)
    print(f"  OK DOCX: {docx_path.name}")
    try:
        d2p_convert(str(docx_path), str(pdf_path))
        print(f"  OK PDF : {pdf_path.name}")
    except Exception as e:
        print(f"  WARN PDF: {e}")
    return docx_path, pdf_path


# ===================================================================
# 1) Ley Reformatoria Segunda Etapa v2 (+ VNE + EPE + excepciones)
# ===================================================================

def gen_ley_etapa_ii_v2():
    print("\n[1] Generando LEY_REFORMATORIA Segunda Etapa v2 ...")
    src = BASE / "02 Segunda Etapa Legislativa" / "LEY_REFORMATORIA_COMYF_LMV_COPLAFIP_Segunda_Etapa_v1.md"
    md = src.read_text(encoding="utf-8")

    # Rebrand global
    md = rebrand(md)

    EPE = """### Artículo 9 bis.- Agrégase el artículo innumerado 165.1 al Código Orgánico Monetario y Financiero. Después del artículo 165, incorpórese el siguiente artículo:

"Artículo 165.1.- Emisores de Pago Electrónico autorizados. La Junta de Política y Regulación Financiera y Monetaria, en coordinación con el Banco Central del Ecuador, autorizará a entidades financieras y a entidades de servicios auxiliares para operar como Emisores de Pago Electrónico de las operaciones liquidadas sobre la Infraestructura de Confianza y Equivalencia Funcional, bajo principio de entrega contra pago en tiempo real. El Emisor de Pago Electrónico opera exclusivamente sobre saldos denominados en dólares de los Estados Unidos de América, sin que tal habilitación constituya emisión de moneda alternativa. Los Emisores de Pago Electrónico cumplirán los requisitos de capital, gobierno corporativo, ciberseguridad y plan de continuidad operativa que apruebe la Junta, así como las obligaciones reforzadas de prevención del lavado de activos previstas en la Ley Orgánica de la materia. La autorización para operar como Emisor de Pago Electrónico no exime de las demás autorizaciones, registros y obligaciones prudenciales que correspondan conforme a este Código."

"""

    VNE = """### Artículo 16 bis.- Refórmase el artículo 26 de la Ley de Mercado de Valores. Agréguese, como inciso final, el siguiente texto:

"Los valores negociables endosables, incluida la factura comercial negociable regulada en el artículo 56 de esta Ley, podrán circular como activos digitales endosables anotados en la Infraestructura de Confianza y Equivalencia Funcional. La cadena de endosos se materializará mediante asientos sucesivos criptográficamente vinculados, sin que la digitalización altere los derechos, obligaciones y excepciones cambiarias propios del título. La presentación, protesto, cobro judicial y demás actos derivados gozarán de plena equivalencia funcional respecto del régimen cartular. La Junta de Política y Regulación Financiera y Monetaria expedirá la norma técnica sobre la interoperabilidad de las plataformas de factoring electrónico que utilicen este régimen."

"""

    EXCEPCIONES_DG = """

SEXTA.- ACTIVOS Y EMISIONES EXCLUIDOS DEL RÉGIMEN. Quedan excluidos del régimen habilitado por la presente Ley:

1. Los valores emitidos por el Estado a través del Ministerio de Economía y Finanzas cuya regulación específica disponga su conservación obligatoria en soporte cartular, sin perjuicio de la representación digital prevista en los artículos 30 y 31 de la presente Ley.

2. Los pasivos directos del Banco Central del Ecuador, cuyo régimen monetario es de competencia exclusiva del organismo emisor conforme al artículo 303 de la Constitución.

3. Las operaciones que no comporten transferencia efectiva de valor o de productos, incluidas las operaciones de aparente compraventa entre cuentas vinculadas o controladas por el mismo beneficiario final con propósito de manipular precios o volúmenes, las cuales constituyen infracción muy grave conforme al artículo 207 reformado de la Ley de Mercado de Valores.

4. La emisión, oferta o intermediación de instrumentos digitales que pretendan operar como moneda de curso legal alternativa al dólar de los Estados Unidos de América.

5. Las emisiones realizadas por personas naturales o jurídicas no inscritas en el Catastro Público del Mercado de Valores, las cuales se consideran oferta pública no autorizada y son sancionadas conforme a la Ley.

6. Los actos enumerados en el artículo 14 bis de la Ley Orgánica de la Fe Pública Digital, en lo que corresponda al mercado de valores.

SÉPTIMA.- FACULTAD DE RESTRICCIÓN. El Directorio del Banco Central del Ecuador y la Junta de Política y Regulación Financiera y Monetaria conservan la facultad de restringir, mediante norma de carácter general motivada en estabilidad financiera, soberanía monetaria o protección al inversionista, el alcance de los activos que pueden representarse como anotación en cuenta sobre la Infraestructura de Confianza y Equivalencia Funcional. La facultad se ejercerá con respeto al debido proceso consagrado en el artículo 76 de la Constitución."""

    md_v2 = md
    md_v2 = md_v2.replace("### Artículo 10.-", EPE + "### Artículo 10.-", 1)
    md_v2 = md_v2.replace("### Artículo 17.-", VNE + "### Artículo 17.-", 1)
    md_v2 = re.sub(
        r"(QUINTA\.\-[^\n]*\n(?:[^\n]*\n)*?)(\n# DISPOSICIONES TRANSITORIAS)",
        r"\1" + EXCEPCIONES_DG + r"\n\2",
        md_v2,
        count=1,
        flags=re.MULTILINE,
    )

    # Etiqueta de version en encabezado
    md_v2 = md_v2.replace(
        "Segunda Etapa Legislativa del programa EcuaLedger Soberana",
        "Segunda Etapa Legislativa del programa EcuaLedger Soberana - VERSIÓN 2\n\n(Integra enmiendas de la Auditoría comparada Ley 7572/25 Paraguay: VNE digital, Emisor de Pago Electrónico, excepciones; renombrado de EcuaBlock a EcuaLedger Soberana)",
        1,
    )

    doc = setup_doc()
    render_block_text(doc, md_v2)
    return save_docx_pdf(doc, "LEY_REFORMATORIA_COMYF_LMV_COPLAFIP_Segunda_Etapa_v2")


# ===================================================================
# 2) Reglamento Segunda Etapa v2
# ===================================================================

def gen_reglamento_etapa_ii_v2():
    print("\n[2] Generando REGLAMENTO Segunda Etapa v2 ...")
    src_docx = BASE / "02 Segunda Etapa Legislativa" / "REGLAMENTO_TECNICO_LEY_REFORMATORIA_Segunda_Etapa_v1.docx"
    src_doc = Document(src_docx)
    md = "\n".join(p.text for p in src_doc.paragraphs)
    md = rebrand(md)

    NUEVO = """

### Artículo 12 bis.- Régimen del Valor Negociable Endosable Digital.

La factura comercial negociable y demás valores endosables anotados en la Infraestructura observarán las siguientes reglas:

1. Cada endoso constituirá asiento sucesivo en la Infraestructura, vinculado al asiento anterior por huella criptográfica, sin solución de continuidad.

2. Los regímenes de protesto, presentación al cobro, aceptación y aval previstos en la Ley Orgánica de Régimen Cambiario, en el Código de Comercio y en la Ley de Mercado de Valores serán plenamente aplicables a la versión digital.

3. La Superintendencia de Compañías, Valores y Seguros llevará un registro de plataformas de factoring electrónico autorizadas e interoperables con la Infraestructura.

4. La sociedad emisora del título original conservará el deber de información al deudor cedido cuando el valor circule electrónicamente.

### Artículo 12 ter.- Régimen del Emisor de Pago Electrónico.

Los Emisores de Pago Electrónico autorizados operarán bajo las siguientes condiciones:

1. Constituirán cuentas de uso exclusivo en el Banco Central del Ecuador para respaldar uno a uno los saldos dispuestos a sus usuarios.

2. Mantendrán capital mínimo, gobierno corporativo, plan de ciberseguridad, póliza de responsabilidad civil profesional y plan de continuidad operativa conforme a las normas que apruebe la Junta de Política y Regulación Financiera y Monetaria.

3. Reportarán a la Unidad de Análisis Financiero y Económico las operaciones inusuales conforme al enfoque basado en riesgo.

4. Publicarán mensualmente en la Infraestructura la auditoría de saldos custodiados con respaldo bancario directo en el Banco Central.

### Artículo 12 quáter.- Activos excluidos del régimen de tokenización.

Conforme a la Disposición General Sexta de la Ley, no podrán tokenizarse sobre la Infraestructura: los pasivos directos del Banco Central, los instrumentos que pretendan funcionar como moneda alternativa al dólar, y las operaciones sin transferencia efectiva de valor. La Superintendencia de Compañías, Valores y Seguros publicará en formato legible por máquina el catálogo dinámico de activos no tokenizables.
"""

    if "Artículo 13.-" in md:
        md_v2 = md.replace("Artículo 13.-", NUEVO + "\nArtículo 13.-", 1)
    else:
        md_v2 = md + "\n\n## Articulado complementario sobre enmiendas\n" + NUEVO

    doc = setup_doc()
    add_titulo(doc, "REGLAMENTO TÉCNICO DE APLICACIÓN DE LA LEY ORGÁNICA REFORMATORIA AL COMYF, LMV Y COPLAFIP", size=14)
    add_p(doc, "Segunda Etapa Legislativa del programa EcuaLedger Soberana - VERSIÓN 2", italic=True)
    add_p(doc, "(Adapta el régimen al Valor Negociable Endosable Digital, al Emisor de Pago Electrónico y a los activos excluidos conforme a la Auditoría Py 7572/25; rebrand EcuaBlock → EcuaLedger Soberana)", italic=True)
    doc.add_paragraph()
    render_block_text(doc, md_v2)
    return save_docx_pdf(doc, "REGLAMENTO_TECNICO_LEY_REFORMATORIA_Segunda_Etapa_v2")


# ===================================================================
# 3) OPCION B - Cap. VII bis LMV v2 (5 enmiendas + excepciones)
# ===================================================================

def gen_opcion_b_v2():
    print("\n[3] Generando OPCION B (Cap VII bis LMV) v2 ...")
    src = BASE / "03 Tercera Etapa Legislativa" / "OPCION_B_Capitulo_Comercializacion_Tokenizada_LMV_v1.md"
    md = src.read_text(encoding="utf-8")
    md = rebrand(md)

    SECAS = """

### Artículo (...)bis 4.A.- Sociedades Emisoras de Capital Abierto Simplificadas.

Reconócese la categoría de Sociedades Emisoras de Capital Abierto Simplificadas (SECAS) como vehículo de emisión de activos digitales representativos de derecho. La SECAS tendrá objeto social exclusivo, capital mínimo, gobierno corporativo digital, política de transparencia mínima y comité de auditoría conforme a la norma que apruebe la Junta de Política y Regulación Financiera y Monetaria. Su régimen es intermedio entre el de la sociedad anónima cerrada y el de la sociedad cuyas acciones se negocien en el segmento principal del mercado de valores. La SECAS podrá emitir hasta los montos máximos por programa que defina la Junta sin requerir prospecto extenso.
"""

    SIMPLIFICADO_MIPYMES = """

### Artículo (...)bis 5.A.- Régimen simplificado de oferta pública digital para MIPYMEs y empresas de base tecnológica.

La Superintendencia de Compañías, Valores y Seguros, mediante norma de carácter general aprobada por la Junta de Política y Regulación Financiera y Monetaria, establecerá un régimen de inscripción simplificada para las emisiones públicas digitales de micro, pequeñas y medianas empresas y de empresas de base tecnológica, hasta los montos máximos por emisión y por colocador que en dicho régimen se determinen. El régimen simplificado exigirá prospecto resumido, calificación de riesgo cuando exceda el umbral de minimis, deber de información continua proporcional al tamaño del emisor y mecanismos de protección reforzada del inversionista minorista, incluida la cuota máxima de exposición prevista en el artículo (...)bis 13.
"""

    BOLSA_PRODUCTOS_FFD = """

## Capítulo III bis - De la Bolsa Soberana de Productos y del Fideicomiso Financiero Digital

### Artículo (...)bis 10.A.- Bolsa Soberana de Productos.

Créase la Bolsa Soberana de Productos del Ecuador, persona jurídica de derecho privado autorizada por la Superintendencia de Compañías, Valores y Seguros, cuyo objeto es organizar la concurrencia de oferta y demanda de bienes básicos representados como activos del mundo real tokenizados sobre la Infraestructura de Confianza y Equivalencia Funcional.

### Artículo (...)bis 10.B.- Productos negociables en la Bolsa Soberana.

Son negociables en la Bolsa Soberana de Productos los activos digitales representativos de cacao, banano, camarón, café, flores, oro, plata, cobre, petróleo, gas, energía eléctrica, créditos de carbono, derechos sobre cosechas, certificados de depósito y warrants electrónicos. La Junta de Política y Regulación Financiera y Monetaria, en coordinación con el Ministerio de Producción, Comercio Exterior, Inversiones y Pesca, autorizará nuevas categorías y los lineamientos de calidad del subyacente.

### Artículo (...)bis 10.C.- Contraparte central temporal.

Hasta que se constituyan cámaras de compensación digitales específicas para productos, la Bolsa Soberana de Productos podrá actuar como contraparte central de liquidación de las operaciones cerradas en su plataforma, conforme a los Principios para Infraestructuras de Mercados Financieros del Comité de Pagos e Infraestructuras de Mercados Financieros y de la Organización Internacional de Comisiones de Valores y al régimen de separación patrimonial que apruebe la Junta de Política y Regulación Financiera y Monetaria.

### Artículo (...)bis 10.D.- Fideicomiso Financiero Digital.

Las sociedades administradoras de fondos y fideicomisos podrán constituir Fideicomisos Financieros Digitales como patrimonios autónomos cuyo activo subyacente sea tokenizable conforme al presente Título. Los certificados de participación, certificados de deuda y certificados subordinados emitidos por el Fideicomiso Financiero Digital tendrán naturaleza de activos digitales representativos de derecho. La Superintendencia de Compañías, Valores y Seguros autorizará el reglamento de gestión y el prospecto de cada Fideicomiso Financiero Digital, con calificación de riesgo obligatoria.
"""

    SEGURO_CIBER = """ El custodio digital cualificado contratará y mantendrá vigente una póliza de seguro de ciberseguridad cuya suma asegurada y cobertura serán determinadas por la Superintendencia de Compañías, Valores y Seguros en proporción al valor de los activos en custodia, sin que dicha suma asegurada pueda ser inferior al diez por ciento del valor agregado de las posiciones custodiadas. La suspensión, no renovación o reducción de la cobertura sin notificación previa a la Superintendencia constituye infracción muy grave."""

    EXCEPCIONES_OB = """

## Capítulo VII bis adicional - Régimen de exclusiones

### Artículo (...)bis 31.- Activos y emisiones excluidos.

Quedan excluidos del régimen de la comercialización tokenizada regulado en el presente Título:

1. Los pasivos directos del Banco Central del Ecuador, en su carácter de instrumentos monetarios de competencia exclusiva del organismo emisor.

2. Las emisiones por personas no inscritas en el Catastro Público del Mercado de Valores, consideradas oferta pública no autorizada.

3. Los instrumentos digitales que pretendan operar como moneda de curso legal alternativa al dólar de los Estados Unidos de América.

4. Las operaciones que no comporten transferencia efectiva de valor o de productos, incluida la simulación de volumen entre cuentas vinculadas.

5. Los actos enumerados en el artículo 14 bis de la Ley Orgánica de la Fe Pública Digital, en lo que corresponda al mercado de valores.

6. Los documentos clasificados como reservados o secretos que afecten a la seguridad del Estado.

### Artículo (...)bis 32.- Facultad de restricción del regulador.

El Directorio del Banco Central del Ecuador y la Junta de Política y Regulación Financiera y Monetaria conservan la facultad de restringir el alcance de los activos representables como anotación en cuenta sobre la Infraestructura, mediante norma de carácter general motivada en estabilidad financiera, soberanía monetaria o protección al inversionista, con respeto al debido proceso del artículo 76 de la Constitución.

### Artículo (...)bis 33.- Conductas tipificadas reforzadas.

Sin perjuicio del artículo (...)bis 27, constituyen infracciones muy graves: a) el uso del término "valor", "inversión" o equivalentes en activos no autorizados por la Superintendencia; b) la entrega de información falsa al regulador respecto del subyacente tokenizado; c) la operación de nodo validador con conocimiento de la ilicitud de las transacciones validadas, hipótesis que excluye la defensa de neutralidad algorítmica del operador.
"""

    md_v2 = md
    md_v2 = md_v2.replace("### Artículo (...)bis 5.-", SECAS + "### Artículo (...)bis 5.-", 1)
    md_v2 = md_v2.replace("### Artículo (...)bis 6.-", SIMPLIFICADO_MIPYMES + "### Artículo (...)bis 6.-", 1)
    md_v2 = md_v2.replace("## Capítulo III —", BOLSA_PRODUCTOS_FFD + "\n## Capítulo III —", 1)
    md_v2 = re.sub(
        r"(### Artículo \(\.\.\.\)bis 17\.- Autorización y requisitos reforzados\.[^\n]*\n[^\n]*continuidad operativa\.)",
        r"\1" + SEGURO_CIBER,
        md_v2,
        count=1,
    )
    if "DISPOSICIONES REFORMATORIAS COMPLEMENTARIAS" in md_v2:
        md_v2 = md_v2.replace(
            "# DISPOSICIONES REFORMATORIAS COMPLEMENTARIAS",
            EXCEPCIONES_OB + "\n# DISPOSICIONES REFORMATORIAS COMPLEMENTARIAS",
            1,
        )
    md_v2 = md_v2.replace(
        "Opción B - Reforma a la Ley de Mercado de Valores",
        "Opción B - Reforma a la Ley de Mercado de Valores - VERSIÓN 2\n\n(Integra enmiendas de la Auditoría Py 7572/25: Bolsa Soberana de Productos, Fideicomiso Financiero Digital, SECAS, régimen simplificado MIPYMEs, seguro ciberseguridad, excepciones; rebrand EcuaBlock → EcuaLedger Soberana)",
        1,
    )

    doc = setup_doc()
    render_block_text(doc, md_v2)
    return save_docx_pdf(doc, "OPCION_B_Capitulo_Comercializacion_Tokenizada_LMV_v2")


# ===================================================================
# 4) OPCION A - LOCATok v2
# ===================================================================

def gen_opcion_a_v2():
    print("\n[4] Generando OPCION A (LOCATok) v2 ...")
    src_docx = BASE / "03 Tercera Etapa Legislativa" / "OPCION_A_LOC-RWA_Ley_Comercializacion_Tokenizada_v1.docx"
    src_doc = Document(src_docx)
    md = "\n".join(p.text for p in src_doc.paragraphs)
    md = rebrand(md)

    APENDICE = """

## Capítulo VII bis (nuevo) - Régimen de la Bolsa Soberana de Productos y del Fideicomiso Financiero Digital

### Artículo 56 bis.- Bolsa Soberana de Productos.

Créase la Bolsa Soberana de Productos del Ecuador, persona jurídica de derecho privado autorizada por la Superintendencia de Compañías, Valores y Seguros, cuyo objeto es organizar la concurrencia de oferta y demanda de bienes básicos representados como activos del mundo real tokenizados sobre la Infraestructura de Confianza y Equivalencia Funcional.

### Artículo 56 ter.- Productos negociables.

Son negociables en la Bolsa Soberana de Productos los activos digitales representativos de cacao, banano, camarón, café, flores, oro, plata, cobre, petróleo, gas, energía eléctrica, créditos de carbono, derechos sobre cosechas, certificados de depósito y warrants electrónicos. La Junta de Política y Regulación Financiera y Monetaria, en coordinación con el Ministerio de Producción, Comercio Exterior, Inversiones y Pesca, autorizará nuevas categorías y los lineamientos de calidad del subyacente.

### Artículo 56 quáter.- Contraparte central temporal.

Hasta que se constituyan cámaras de compensación digitales específicas para productos, la Bolsa Soberana de Productos podrá actuar como contraparte central de liquidación de las operaciones cerradas en su plataforma, conforme a los Principios para Infraestructuras de Mercados Financieros y al régimen de separación patrimonial que apruebe la Junta de Política y Regulación Financiera y Monetaria.

### Artículo 56 quinquies.- Fideicomiso Financiero Digital.

Las sociedades administradoras de fondos y fideicomisos podrán constituir Fideicomisos Financieros Digitales como patrimonios autónomos cuyo activo subyacente sea tokenizable conforme a la presente Ley. Los certificados de participación, certificados de deuda y certificados subordinados emitidos por el Fideicomiso Financiero Digital tendrán naturaleza de activos digitales representativos de derecho. La Superintendencia de Compañías, Valores y Seguros autorizará el reglamento de gestión y el prospecto de cada Fideicomiso Financiero Digital, con calificación de riesgo obligatoria.

### Artículo 56 sexies.- Sociedades Emisoras de Capital Abierto Simplificadas.

Reconócese la categoría de Sociedades Emisoras de Capital Abierto Simplificadas (SECAS) como vehículo de emisión de activos digitales representativos de derecho. La SECAS tendrá objeto social exclusivo, capital mínimo, gobierno corporativo digital, política de transparencia mínima y comité de auditoría conforme a la norma que apruebe la Junta de Política y Regulación Financiera y Monetaria.

### Artículo 56 septies.- Régimen simplificado para MIPYMEs.

La Superintendencia de Compañías, Valores y Seguros, mediante norma de carácter general aprobada por la Junta de Política y Regulación Financiera y Monetaria, establecerá un régimen de inscripción simplificada para las emisiones públicas digitales de micro, pequeñas y medianas empresas y de empresas de base tecnológica, con prospecto resumido y deber de información proporcional al tamaño del emisor.

### Artículo 56 octies.- Seguro de ciberseguridad obligatorio del custodio digital.

El custodio digital cualificado contratará y mantendrá vigente una póliza de seguro de ciberseguridad cuya suma asegurada y cobertura serán determinadas por la Superintendencia de Compañías, Valores y Seguros en proporción al valor de los activos en custodia, sin que dicha suma asegurada pueda ser inferior al diez por ciento del valor agregado de las posiciones custodiadas. La suspensión, no renovación o reducción de la cobertura sin notificación previa a la Superintendencia constituye infracción muy grave.

## Capítulo VII ter (nuevo) - Régimen de exclusiones

### Artículo 56 nonies.- Activos y emisiones excluidos.

Quedan excluidos del régimen de la presente Ley:

1. Los pasivos directos del Banco Central del Ecuador.
2. Las emisiones por personas no inscritas en el Catastro Público del Mercado de Valores.
3. Los instrumentos digitales que pretendan operar como moneda alternativa al dólar de los Estados Unidos de América.
4. Las operaciones que no comporten transferencia efectiva de valor o de productos.
5. Los actos enumerados en el artículo 14 bis de la Ley Orgánica de la Fe Pública Digital.
6. Los documentos clasificados como reservados o secretos.

### Artículo 56 decies.- Facultad de restricción.

El Directorio del Banco Central del Ecuador y la Junta de Política y Regulación Financiera y Monetaria conservan la facultad de restringir, mediante norma de carácter general motivada, el alcance de los activos tokenizables, con respeto al debido proceso del artículo 76 de la Constitución.

### Artículo 56 undecies.- Conductas tipificadas reforzadas.

Constituyen infracciones muy graves: a) el uso del término "valor", "inversión" o equivalentes en activos no autorizados; b) la entrega de información falsa al regulador respecto del subyacente tokenizado; c) la operación de nodo validador con conocimiento de la ilicitud de las transacciones validadas, sin posibilidad de invocar neutralidad algorítmica.
"""
    if "DISPOSICIONES GENERALES" in md:
        md_v2 = md.replace("DISPOSICIONES GENERALES", APENDICE + "\nDISPOSICIONES GENERALES", 1)
    elif "DISPOSICIONES TRANSITORIAS" in md:
        md_v2 = md.replace("DISPOSICIONES TRANSITORIAS", APENDICE + "\nDISPOSICIONES TRANSITORIAS", 1)
    else:
        md_v2 = md + APENDICE

    doc = setup_doc()
    add_titulo(doc, "PROYECTO DE LEY ORGÁNICA DE COMERCIALIZACIÓN TOKENIZADA DE ACTIVOS DEL MUNDO REAL (LOC-RWA)", size=14)
    add_p(doc, "Tercera Etapa Legislativa del programa EcuaLedger Soberana - Opción A - VERSIÓN 2", italic=True)
    add_p(doc, "(Integra enmiendas de la Auditoría Py 7572/25: Bolsa Soberana de Productos, FFD, SECAS, MIPYMEs, seguro ciberseguridad, excepciones; rebrand EcuaBlock → EcuaLedger Soberana)", italic=True)
    doc.add_paragraph()
    render_block_text(doc, md_v2)
    return save_docx_pdf(doc, "OPCION_A_LOC-RWA_Ley_Comercializacion_Tokenizada_v2")


# ===================================================================
# 5) Reglamento Tercera Etapa v2 (adapta ambas opciones)
# ===================================================================

def gen_reglamento_etapa_iii_v2():
    print("\n[5] Generando REGLAMENTO Tercera Etapa v2 ...")
    src_docx = BASE / "03 Tercera Etapa Legislativa" / "REGLAMENTO_TECNICO_LEY_COMERCIALIZACION_TOKENIZADA_Tercera_Etapa_v1.docx"
    src_doc = Document(src_docx)
    md = "\n".join(p.text for p in src_doc.paragraphs)
    md = rebrand(md)

    NUEVO = """

### Artículo 24 bis.- Operación de la Bolsa Soberana de Productos.

La Bolsa Soberana de Productos operará bajo:

1. Gobernanza tripartita en proporciones equivalentes Estado-sector privado-academia.
2. Sistema de calidad y trazabilidad del subyacente, con certificación de terceros independientes.
3. Sistema de calificación crediticia de contrapartes.
4. Garantías ejecutables sobre el subyacente y sobre cuentas de margen.
5. Reglamento interno disciplinario aprobado por la Superintendencia de Compañías, Valores y Seguros.

### Artículo 24 ter.- Categorías iniciales de productos tokenizables.

Categorías habilitadas desde el inicio: cacao CCN-51 y nacional fino de aroma; banano por variedad; camarón con trazabilidad de origen; café arábica y robusta; flores frescas de exportación; oro Au999 y Au995; plata; cobre; petróleo crudo Oriente y Napo; gas natural; energía hidroeléctrica; créditos de carbono certificados; derechos sobre cosechas anuales; warrants electrónicos sobre almacenes generales autorizados.

### Artículo 24 quáter.- Régimen del Fideicomiso Financiero Digital.

El Fideicomiso Financiero Digital observará:

1. Patrimonio autónomo segregado, registrado en la Infraestructura.
2. Tres clases de certificados: participación, deuda y subordinados.
3. Calificación de riesgo obligatoria para emisiones colocadas mediante oferta pública.
4. Reglamento de gestión publicado en la Infraestructura.
5. Auditoría externa permanente.

### Artículo 24 quinquies.- SECAS - Sociedades Emisoras de Capital Abierto Simplificadas.

Las SECAS observarán: capital mínimo de cien mil dólares; comité de auditoría con al menos un miembro independiente; política de divulgación periódica; gobierno corporativo digital con asambleas mediante firma electrónica cualificada; cumplimiento del régimen de mercado abierto en lo que corresponda al volumen de emisión.

### Artículo 24 sexies.- Régimen simplificado MIPYMEs.

El régimen simplificado para MIPYMEs y empresas de base tecnológica observará: prospecto resumido de hasta veinte páginas; calificación de riesgo facultativa para emisiones de hasta quinientos mil dólares; reportes anuales en formato simplificado; cuota máxima de exposición del inversionista minorista de cinco mil dólares por emisión sin asesoramiento profesional.

### Artículo 24 septies.- Seguro de ciberseguridad del custodio digital.

El seguro de ciberseguridad observará: suma asegurada mínima del diez por ciento del valor agregado de posiciones custodiadas; cobertura de eventos de ciberataque, ransomware, ingeniería social, error humano y compromiso de claves; aseguradora con calificación de riesgo internacional grado de inversión; vigencia anual renovable; notificación previa a la Superintendencia de cualquier modificación.

### Artículo 24 octies.- Catálogo dinámico de activos excluidos.

La Superintendencia de Compañías, Valores y Seguros mantendrá un catálogo dinámico de los activos y emisiones excluidos del régimen tokenizado, conforme a las exclusiones previstas en la Ley. El catálogo se publicará en la Infraestructura, en formato legible por máquina, con actualización semestral mínima.

### Artículo 24 nonies.- Régimen del Emisor de Pago Electrónico para liquidación.

Cuando las operaciones tokenizadas se liquiden mediante Emisor de Pago Electrónico autorizado conforme a la Ley Reformatoria Segunda Etapa, observarán las reglas operativas, prudenciales y de prevención del lavado de activos contenidas en el reglamento técnico de la Segunda Etapa.

### Artículo 24 decies.- Articulación con el régimen del Valor Negociable Endosable.

Los valores negociables endosables tokenizados conforme al artículo 16 bis de la Ley Reformatoria Segunda Etapa podrán negociarse en los mercados secundarios digitales regulados en el presente reglamento, siempre que las plataformas estén inscritas en el Catastro Público del Mercado de Valores y observen el principio de mejor ejecución.
"""

    if "Artículo 25.-" in md:
        md_v2 = md.replace("Artículo 25.-", NUEVO + "\nArtículo 25.-", 1)
    else:
        md_v2 = md + "\n\n## Articulado complementario sobre enmiendas\n" + NUEVO

    doc = setup_doc()
    add_titulo(doc, "REGLAMENTO TÉCNICO DE APLICACIÓN DE LA LEY DE COMERCIALIZACIÓN TOKENIZADA", size=14)
    add_p(doc, "Tercera Etapa Legislativa del programa EcuaLedger Soberana - VERSIÓN 2", italic=True)
    add_p(doc, "(Adapta el régimen a la Bolsa Soberana de Productos, FFD, SECAS, MIPYMEs, seguro ciberseguridad, EPE, VNE y excepciones de la Auditoría Py 7572/25; rebrand EcuaBlock → EcuaLedger Soberana)", italic=True)
    doc.add_paragraph()
    render_block_text(doc, md_v2)
    return save_docx_pdf(doc, "REGLAMENTO_TECNICO_LEY_COMERCIALIZACION_TOKENIZADA_Tercera_Etapa_v2")


# ===================================================================
# ORQUESTA
# ===================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("REELABORACION INTEGRAL v2 - ETAPAS II y III")
    print("LOFPD (Etapa I) NO se toca - el usuario la maneja manualmente")
    print("Rebrand: EcuaBlock -> EcuaLedger Soberana")
    print("=" * 60)
    results = []
    for fn in [gen_ley_etapa_ii_v2, gen_reglamento_etapa_ii_v2,
               gen_opcion_b_v2, gen_opcion_a_v2,
               gen_reglamento_etapa_iii_v2]:
        try:
            r = fn()
            results.append(("OK", fn.__name__))
        except Exception as e:
            results.append(("FAIL", f"{fn.__name__}: {e}"))
            import traceback
            traceback.print_exc()
    print("\n" + "=" * 60)
    print("RESUMEN")
    print("=" * 60)
    for st, label in results:
        print(f"  {st:4s}  {label}")
