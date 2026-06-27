# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "E:/vars/var 8/3 Analisis Financiero QVP/Mensaje_WhatsApp_Don_Omar_Oro_Rojo.docx"
NAVY = RGBColor(0x1B,0x2A,0x4A); GRAY = RGBColor(0x6b,0x74,0x80); RED = RGBColor(0xB0,0x24,0x18)

doc = Document()
# margenes
for s in doc.sections:
    s.left_margin=s.right_margin=Inches(0.8); s.top_margin=s.bottom_margin=Inches(0.7)

# Titulo
t=doc.add_paragraph(); r=t.add_run("Mensaje para WhatsApp — Don Omar Juez Zambrano")
r.bold=True; r.font.size=Pt(15); r.font.color.rgb=NAVY
s=doc.add_paragraph(); rr=s.add_run("Hallazgo: préstamo a parte relacionada Oro Rojo / OroJuez S.A.")
rr.font.size=Pt(10); rr.font.color.rgb=GRAY

# Instruccion
ins=doc.add_paragraph()
ri=ins.add_run("Instrucciones: copie únicamente el texto encuadrado de abajo y péguelo en el chat de WhatsApp. "
 "Los *asteriscos* se convierten automáticamente en negrita dentro de WhatsApp; no los borre.")
ri.font.size=Pt(9); ri.italic=True; ri.font.color.rgb=RED
doc.add_paragraph("─"*60).runs[0].font.color.rgb=GRAY

# ===== MENSAJE (texto literal con sintaxis WhatsApp) =====
msg_lines = [
 "*RESUMEN — Préstamo a Oro Rojo / OroJuez S.A.*",
 "",
 "Don Omar, buenos días.",
 "",
 "Terminé el análisis financiero de Quevepalma (2024–2025) y revisé en detalle el mayor contable del préstamo a la relacionada Oro Rojo / OroJuez. Le resumo lo más importante:",
 "",
 "💰 *El préstamo*",
 "• Capital: $1.551.059 (may-2023) + $150.000 de interés único = *$1.701.059*",
 "• Cuota fija: $28.351/mes a 60 meses (madura may-2028)",
 "• Saldo hoy: *$680.424* (60% pagado)",
 "",
 "📉 *El punto clave*",
 "La tasa efectiva de ese préstamo es *3,76% anual*. Pero Quevepalma se endeuda con los bancos (CFN, Pacífico, Produbanco) al *~10%*. Esa diferencia de *6,24 puntos* es un subsidio implícito a la relacionada.",
 "",
 "En palabras simples: estamos pidiendo caro a los bancos para prestar barato a Oro Rojo. Cada año, la diferencia sale del bolsillo de la empresa.",
 "",
 "💵 *Cuánto nos cuesta*",
 "El costo de oportunidad acumulado es *~$265.513* en la vida del préstamo. Para dimensionarlo: equivale a *~20% de la pérdida de $1,3 millones* que tuvo Quevepalma en 2025.",
 "",
 "⚠️ *Dato adicional*",
 "El expediente incluye una línea \"DEUDA DE EPACEM\" de $226.832 que conecta directamente con el caso EPACEM que ya venimos revisando.",
 "",
 "📎 Le preparé un informe ejecutivo de 2 páginas (con un anexo que explica el cálculo paso a paso) y el modelo financiero completo en Excel. Se los adjunto a continuación.",
 "",
 "Quedo atento a sus comentarios.",
]
for line in msg_lines:
    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(2); p.paragraph_format.space_before=Pt(0)
    rn=p.add_run(line); rn.font.size=Pt(11); rn.font.name="Calibri"

doc.add_paragraph("─"*60).runs[0].font.color.rgb=GRAY

# Nota sobre adjuntos
n=doc.add_paragraph()
rn=n.add_run("Nota sobre el adjunto: para WhatsApp en el celular, el PDF "
 "(INFORME_FORENSE_ORO_ROJO_Directorio.pdf) se ve mejor que el HTML — se abre directo como documento. "
 "El HTML conviene para correo o para editar. Ambos están en la carpeta del análisis.")
rn.font.size=Pt(9); rn.italic=True; rn.font.color.rgb=GRAY

doc.save(OUT)
import os
print("DOCX generado:", OUT)
print("Tamaño:", os.path.getsize(OUT), "bytes")
