# -*- coding: utf-8 -*-
"""
INFORME_REFORMA_COMYF_IBPP_2026_v4_LOFPD - Guía de Implementación
==================================================================
- Reformula v3 como guía explicativa de estrategia de implementación
- Integra enmiendas v2 de Etapas II y III (auditoría Py 7572/25)
- Elimina referencias nominales a la Asambleísta; conserva solo "la Comisión"
- Inteligencia política: sutil, velada, no manifestada
- Rebrand: EcuaBlock → EcuaLedger Soberana (ya presente en v3)
- Output: DOCX + PDF (no MD)
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert as d2p_convert

BASE = Path(r"C:\Users\datos\Dropbox\var 91\Neg Inm\1 Proy Belen 2026"
            r"\Tokenizacion Activos Reales - RWA\JRPFM"
            r"\Proyecto Marco Regulatorio Asamblea")
DEST = BASE / "00b Docs Legales"

PURPURA = RGBColor(0x72, 0x2F, 0xA1)


def setup_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Cambria"
    style.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    return doc


def h1(doc, txt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = PURPURA


def h2(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = PURPURA


def h3(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(11.5)


def p(doc, txt, italic=False, justify=True):
    par = doc.add_paragraph()
    if justify:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = par.add_run(txt)
    if italic:
        r.italic = True


def bul(doc, txt):
    doc.add_paragraph(txt, style="List Bullet")


def num(doc, txt):
    doc.add_paragraph(txt, style="List Number")


def portada(doc):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("REPÚBLICA DEL ECUADOR")
    r1.bold = True
    r1.font.size = Pt(13)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("ASAMBLEA NACIONAL — COMISIÓN ESPECIALIZADA PERMANENTE DE RÉGIMEN ECONÓMICO Y TRIBUTARIO")
    r2.bold = True
    r2.font.size = Pt(11)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "GUÍA DE IMPLEMENTACIÓN JURÍDICA DEL PROGRAMA ECUALEDGER SOBERANA"
    )
    r3.bold = True
    r3.font.size = Pt(16)
    r3.font.color.rgb = PURPURA

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(
        "Estrategia legislativa en tres etapas para la consolidación de la "
        "Infraestructura Blockchain Pública Permisionada (IBPP) del Ecuador, "
        "con integración de las enmiendas derivadas de la auditoría comparada a la "
        "Ley N.° 7572/25 del Paraguay"
    )
    r4.italic = True
    r4.font.size = Pt(11)

    doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("Versión 4 — junio de 2026")
    r5.bold = True
    r5.font.size = Pt(11)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run(
        "Sustituye a la versión 3 del Informe de Contribuciones Jurídico-Constitucionales. "
        "Incorpora las enmiendas v2 de Segunda y Tercera Etapa Legislativa derivadas de la "
        "auditoría de simetría normativa con la Ley N.° 7572/25 del Paraguay, así como la "
        "armonización terminológica del programa con el nombre EcuaLedger Soberana."
    )
    r6.italic = True
    r6.font.size = Pt(10)

    doc.add_paragraph()
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r7 = p7.add_run("Quito, Distrito Metropolitano")
    r7.font.size = Pt(11)


def build():
    doc = setup_doc()
    portada(doc)
    doc.add_page_break()

    # ====== I. RESUMEN EJECUTIVO ======
    h1(doc, "I. Resumen ejecutivo")
    p(doc,
      "La presente guía expone, de manera ordenada y explicativa, la estrategia "
      "jurídica de implementación del programa EcuaLedger Soberana en el "
      "ordenamiento ecuatoriano. La estrategia se articula en tres etapas "
      "legislativas que se preparan recíprocamente: la Ley Orgánica de la Fe "
      "Pública Digital (Primera Etapa), la Ley Orgánica Reformatoria al Código "
      "Orgánico Monetario y Financiero, a la Ley de Mercado de Valores y al "
      "Código Orgánico de Planificación y Finanzas Públicas (Segunda Etapa) y "
      "el régimen de Comercialización Tokenizada, sea en su Opción A — Ley "
      "Orgánica autónoma — o en su Opción B — Capítulo VII bis incorporado a la "
      "Ley de Mercado de Valores (Tercera Etapa).")
    p(doc,
      "Cada etapa cuenta con su respectivo reglamento técnico de aplicación. "
      "La presente versión 4 incorpora las enmiendas derivadas de la auditoría "
      "comparada con la Ley N.° 7572/25 de la República del Paraguay, que "
      "permiten alcanzar paridad sistémica con el referente regional sin "
      "alterar la arquitectura constitucional ecuatoriana.")
    p(doc,
      "El documento se concibe como guía de orientación para los actores que "
      "intervienen en la implementación: la Comisión Especializada Permanente "
      "de Régimen Económico y Tributario, las autoridades de regulación y "
      "control del sistema financiero y del mercado de valores, los órganos de "
      "coordinación interinstitucional, y las contrapartes técnico-académicas "
      "del programa.")

    # ====== II. MARCO CONSTITUCIONAL ======
    h1(doc, "II. Marco constitucional habilitante")
    p(doc,
      "Las tres etapas legislativas no requieren reforma constitucional. Su "
      "fundamento expreso se encuentra en los siguientes preceptos de la "
      "Constitución de la República:")
    bul(doc, "Arts. 16, num. 2, y 18 — derecho al acceso universal a las TIC e información pública.")
    bul(doc, "Arts. 66, num. 19 y num. 21 — protección de datos personales y secreto de la correspondencia virtual.")
    bul(doc, "Arts. 76 y 82 — debido proceso y seguridad jurídica.")
    bul(doc, "Art. 169 — simplificación, eficacia, celeridad e inmediación procesal.")
    bul(doc, "Arts. 275, 283 y 284 — Régimen de Desarrollo y Sistema Económico Social y Solidario.")
    bul(doc, "Arts. 285-292 — política fiscal, endeudamiento público y rendición de cuentas.")
    bul(doc, "Arts. 302 y 303 — facultad exclusiva del Estado en la política monetaria a través del Banco Central del Ecuador.")
    bul(doc, "Arts. 308, 309 y 310 — actividades financieras como servicio de orden público.")
    bul(doc, "Arts. 313, 314 y 315 — sectores estratégicos y empresas públicas.")
    bul(doc, "Arts. 321 y 322 — formas de propiedad y propiedad intelectual.")
    bul(doc, "Arts. 385-387 — Sistema Nacional de Ciencia, Tecnología, Innovación y Saberes Ancestrales.")
    bul(doc, "Arts. 425 y 438, num. 3 — jerarquía normativa y consulta facultativa de constitucionalidad.")

    # ====== III. MODELO DOCTRINAL PY ======
    h1(doc, "III. Modelo doctrinal paraguayo de referencia")
    p(doc,
      "La estrategia toma como referente comparado la arquitectura jurídica "
      "del Paraguay, estructurada en dos leyes en cascada y un sandbox "
      "regulatorio. Sus principios rectores se trasladan al Ecuador con las "
      "adaptaciones constitucionales correspondientes:")
    bul(doc, "Equivalencia funcional plena entre soporte físico y soporte electrónico (Art. 3 Ley 6822/21).")
    bul(doc, "Documento Transmisible Electrónico (DTE) como instituto medular.")
    bul(doc, "Anotación en cuenta sobre tecnologías de registros distribuidos (Art. 65 Ley 7572/25).")
    bul(doc, "Gobernanza tripartita 33-33-33 — Estado, sector privado y academia.")
    bul(doc, "Soberanía territorial de los nodos validadores y personalidad jurídica nacional verificable.")
    bul(doc, "Liquidación bruta en tiempo real con principio de entrega contra pago, conforme estándares IOSCO-CPMI.")
    p(doc,
      "El presente programa adopta esos principios con las precisiones "
      "específicas del bloque de constitucionalidad ecuatoriano: facultad "
      "exclusiva del Banco Central del Ecuador en política monetaria (Art. 303 "
      "CRE), prohibición de moneda alternativa al dólar de los Estados Unidos "
      "de América y reserva del registro público como servicio de orden público "
      "(Arts. 308 y 313 CRE).")

    h2(doc, "III bis. La Fundación Blockchain Soberana del Ecuador como articuladora técnico-institucional")
    p(doc,
      "La Fundación Blockchain Soberana del Ecuador (FBSE), con sede en Quevedo "
      "(Los Ríos) y constituida el 16 de marzo de 2026, es la entidad sin fines "
      "de lucro encargada de articular técnicamente la infraestructura, "
      "coordinar la gobernanza tripartita y mantener el convenio de "
      "interoperabilidad con la red soberana del Paraguay. Su naturaleza no "
      "lucrativa y su gobernanza compartida con el sector privado y la academia "
      "satisfacen el mandato del Art. 313 CRE sobre la prevención de la "
      "captura regulatoria.")

    h2(doc, "III ter. Bloque transversal de protección de datos personales e inmutabilidad")
    p(doc,
      "La principal objeción que se formula a las infraestructuras de registros "
      "distribuidos es la presunta colisión entre la inmutabilidad técnica del "
      "registro y los derechos consagrados en los arts. 66, num. 19 CRE, y en la "
      "Ley Orgánica de Protección de Datos Personales (LOPDP). El programa "
      "resuelve esta tensión mediante ocho reglas técnico-jurídicas:")
    num(doc, "Off-chaining estricto de datos personales identificables — solo se inscribe la huella criptográfica.")
    num(doc, "Separación de capas (twin-chain) — capa jurídica cifrada extremo a extremo y capa de negocios para activos digitales.")
    num(doc, "Derecho de rectificación ejercido sobre la base de datos off-chain con reflejo de huella en la cadena.")
    num(doc, "Derecho al olvido y a la supresión mediante crypto-shredding de la clave de descifrado.")
    num(doc, "Principio de minimización del dato anotado.")
    num(doc, "Auditoría independiente y derecho al recurso efectivo ante la autoridad de protección de datos.")
    num(doc, "Sello digital de Estado equiparable a la correspondencia virtual del Art. 66, num. 21 CRE.")
    num(doc, "Prohibición expresa de anotar datos sensibles, aun en forma de hash.")

    # ====== IV. LAS TRES ETAPAS ======
    h1(doc, "IV. Las tres etapas legislativas — mapa de la estrategia")
    p(doc,
      "La implementación del programa se articula en tres etapas legislativas "
      "que se preparan recíprocamente. Cada etapa responde a una pregunta "
      "jurídica concreta y prepara el sustrato de la siguiente. La omisión o "
      "alteración del orden vicia las etapas posteriores:")
    h3(doc, "Primera Etapa — Ley Orgánica de la Fe Pública Digital (LOFPD)")
    p(doc,
      "Reconoce la equivalencia funcional plena entre el documento, la firma, "
      "el sello y la anotación electrónicos respecto de sus contrapartes en "
      "soporte físico. Configura la Infraestructura de Confianza y "
      "Equivalencia Funcional, regula el Documento Transmisible Electrónico y "
      "establece el sandbox regulatorio. Pregunta jurídica: ¿el bit equivale al "
      "papel ante la Ley?")
    h3(doc, "Segunda Etapa — Reforma al COMYF, LMV y COPLAFIP")
    p(doc,
      "Habilita las anotaciones en cuenta sobre infraestructura distribuida, "
      "los depósitos centralizados digitales, las casas de valores digitales, "
      "la liquidación bruta en tiempo real, la entrega contra pago en tiempo "
      "real y la tokenización de activos del mundo real con respeto al límite "
      "constitucional del Art. 303 CRE. Pregunta jurídica: ¿el derecho ya en "
      "bit puede ser token?")
    h3(doc, "Tercera Etapa — Régimen de Comercialización Tokenizada")
    p(doc,
      "Regula los mercados secundarios digitales, los intermediarios digitales "
      "cualificados, los custodios digitales, las cámaras de compensación "
      "digitales y la protección reforzada del inversionista digital. Opción A "
      "— ley orgánica autónoma. Opción B — Capítulo VII bis incorporado a la "
      "Ley de Mercado de Valores. Pregunta jurídica: ¿cómo se transa, compensa "
      "y liquida el token en territorio soberano?")

    # ====== V. ETAPA I ======
    h1(doc, "V. Etapa I — Ley Orgánica de la Fe Pública Digital y su Reglamento")
    p(doc,
      "La Ley Orgánica de la Fe Pública Digital, junto con su Reglamento "
      "Técnico, constituye la planta baja del edificio jurídico del programa. "
      "Su contenido sustantivo se mantiene conforme a las versiones vigentes en "
      "el expediente, sin perjuicio de las precisiones técnico-administrativas "
      "que la Función Ejecutiva incorpore en el trámite parlamentario.")
    p(doc,
      "Los siete títulos de la Ley desarrollan: las disposiciones generales, "
      "la equivalencia funcional y los servicios cualificados de confianza, el "
      "Documento Transmisible Electrónico, la Infraestructura de Confianza y "
      "Equivalencia Funcional, el sandbox regulatorio, la Autoridad de "
      "Servicios de Confianza y el régimen sancionatorio. El bloque transversal "
      "de protección de datos personales atraviesa todos los títulos.")

    # ====== VI. ETAPA II CON ENMIENDAS ======
    h1(doc, "VI. Etapa II — Reforma al COMYF, LMV y COPLAFIP, con enmiendas de la auditoría Py 7572/25")
    p(doc,
      "La versión 2 del proyecto de la Segunda Etapa incorpora dos enmiendas "
      "estructurales y un régimen de exclusiones derivados de la auditoría de "
      "simetría normativa con la Ley N.° 7572/25 del Paraguay. La incorporación "
      "es quirúrgica y preserva la unidad codificada del marco financiero "
      "ecuatoriano.")

    h2(doc, "VI.1 Enmienda B — Valor Negociable Endosable Digital")
    p(doc,
      "Se incorpora el régimen de circulación electrónica de los valores "
      "negociables endosables, incluida la factura comercial negociable "
      "regulada en el Art. 56 de la Ley de Mercado de Valores. La cadena de "
      "endosos se materializa mediante asientos sucesivos criptográficamente "
      "vinculados, sin alteración de los derechos, obligaciones y excepciones "
      "cambiarias propios del título. La presentación, el protesto, el cobro "
      "judicial y los demás actos derivados gozan de plena equivalencia "
      "funcional respecto del régimen cartular. Esta enmienda habilita el "
      "factoring electrónico interoperable con la infraestructura del programa.")

    h2(doc, "VI.2 Enmienda C — Emisor de Pago Electrónico")
    p(doc,
      "Se incorpora la figura del Emisor de Pago Electrónico, autorizado por "
      "la Junta de Política y Regulación Financiera y Monetaria en coordinación "
      "con el Banco Central del Ecuador, para liquidar el componente "
      "dinerario de las operaciones tokenizadas bajo principio de entrega "
      "contra pago en tiempo real. El Emisor de Pago Electrónico opera "
      "exclusivamente sobre saldos denominados en dólares de los Estados Unidos "
      "de América, con lo cual se preserva la facultad exclusiva del Banco "
      "Central del Ecuador en materia monetaria (Art. 303 CRE). La figura se "
      "sujeta a requisitos reforzados de capital, gobierno corporativo, "
      "ciberseguridad, continuidad operativa y obligaciones antilavado.")

    h2(doc, "VI.3 Régimen de exclusiones derivado de la Ley N.° 7572/25 del Paraguay")
    p(doc,
      "Se incorporan dos disposiciones generales que delimitan el perímetro "
      "del régimen tokenizado:")
    num(doc, "Disposición General Sexta — Activos y emisiones excluidos del régimen: los pasivos directos del Banco Central del Ecuador, los instrumentos digitales que pretendan operar como moneda alternativa al dólar, las operaciones sin transferencia efectiva de valor, las emisiones por personas no inscritas en el Catastro Público del Mercado de Valores, y los actos enumerados en el Art. 14 bis de la LOFPD.")
    num(doc, "Disposición General Séptima — Facultad de restricción: el Directorio del Banco Central del Ecuador y la Junta de Política y Regulación Financiera y Monetaria conservan la facultad de restringir, mediante norma motivada en estabilidad financiera, soberanía monetaria o protección al inversionista, el alcance de los activos tokenizables.")

    h2(doc, "VI.4 Reglamento Técnico de la Segunda Etapa")
    p(doc,
      "El reglamento de aplicación desarrolla los procedimientos operativos "
      "de las dos enmiendas y del régimen de exclusiones: el registro de "
      "plataformas de factoring electrónico autorizadas, las condiciones de "
      "operación del Emisor de Pago Electrónico (incluida la cuenta de uso "
      "exclusivo en el Banco Central del Ecuador y la auditoría mensual de "
      "saldos), y el catálogo dinámico de activos no tokenizables publicado en "
      "formato legible por máquina.")

    # ====== VII. ETAPA III CON ENMIENDAS ======
    h1(doc, "VII. Etapa III — Comercialización tokenizada, con enmiendas de la auditoría Py 7572/25")
    p(doc,
      "La Tercera Etapa puede materializarse mediante una Ley Orgánica "
      "autónoma (Opción A — LOC-RWA) o mediante un Capítulo VII bis "
      "incorporado a la Ley de Mercado de Valores (Opción B). Ambas opciones "
      "han sido elaboradas en versión 2 con la integración de cinco enmiendas "
      "estructurales y un régimen de exclusiones derivados de la auditoría "
      "comparada con la Ley N.° 7572/25.")

    h2(doc, "VII.1 Enmienda A — Bolsa Soberana de Productos")
    p(doc,
      "Créase la Bolsa Soberana de Productos del Ecuador como persona jurídica "
      "de derecho privado autorizada por la Superintendencia de Compañías, "
      "Valores y Seguros, con el objeto de organizar la concurrencia de oferta "
      "y demanda de bienes básicos representados como activos del mundo real "
      "tokenizados. Son negociables, sin agotar el catálogo: cacao, banano, "
      "camarón, café, flores, oro, plata, cobre, petróleo, gas, energía "
      "eléctrica, créditos de carbono, derechos sobre cosechas, certificados "
      "de depósito y warrants electrónicos. La Bolsa Soberana de Productos "
      "podrá actuar como contraparte central temporal hasta que se constituyan "
      "cámaras de compensación digitales específicas.")

    h2(doc, "VII.2 Enmienda D — Fideicomiso Financiero Digital")
    p(doc,
      "Las sociedades administradoras de fondos y fideicomisos podrán "
      "constituir Fideicomisos Financieros Digitales como patrimonios "
      "autónomos cuyo activo subyacente sea tokenizable. Los certificados de "
      "participación, los certificados de deuda y los certificados "
      "subordinados emitidos por el Fideicomiso Financiero Digital tendrán "
      "naturaleza de activos digitales representativos de derecho. La "
      "Superintendencia de Compañías, Valores y Seguros autoriza el reglamento "
      "de gestión y el prospecto de cada Fideicomiso, con calificación de "
      "riesgo obligatoria.")

    h2(doc, "VII.3 Enmienda G — Sociedades Emisoras de Capital Abierto Simplificadas (SECAS)")
    p(doc,
      "Se reconoce la categoría de Sociedades Emisoras de Capital Abierto "
      "Simplificadas como vehículo de emisión de activos digitales "
      "representativos de derecho. La SECAS tiene objeto social exclusivo, "
      "capital mínimo, gobierno corporativo digital, política de transparencia "
      "mínima y comité de auditoría conforme a la norma que apruebe la Junta "
      "de Política y Regulación Financiera y Monetaria. Su régimen es "
      "intermedio entre la sociedad anónima cerrada y la sociedad cotizada en "
      "el segmento principal del mercado.")

    h2(doc, "VII.4 Enmienda F — Régimen simplificado para MIPYMEs y empresas de base tecnológica")
    p(doc,
      "La Superintendencia de Compañías, Valores y Seguros, mediante norma de "
      "carácter general aprobada por la Junta, establecerá un régimen de "
      "inscripción simplificada para las emisiones públicas digitales de "
      "micro, pequeñas y medianas empresas y de empresas de base tecnológica. "
      "El régimen exige prospecto resumido, calificación de riesgo cuando "
      "exceda el umbral de minimis y deber de información proporcional al "
      "tamaño del emisor.")

    h2(doc, "VII.5 Enmienda H — Seguro de ciberseguridad obligatorio del custodio digital")
    p(doc,
      "El custodio digital cualificado contratará y mantendrá vigente una "
      "póliza de seguro de ciberseguridad cuya suma asegurada y cobertura "
      "serán determinadas por la Superintendencia en proporción al valor de "
      "los activos en custodia, sin que la suma pueda ser inferior al diez por "
      "ciento del valor agregado de las posiciones custodiadas. La "
      "suspensión, no renovación o reducción de la cobertura sin notificación "
      "previa constituye infracción muy grave.")

    h2(doc, "VII.6 Régimen de exclusiones de la Tercera Etapa")
    p(doc,
      "El nuevo Capítulo de exclusiones incorpora: los pasivos directos del "
      "Banco Central del Ecuador, las emisiones por personas no inscritas en "
      "el Catastro Público, los instrumentos digitales que pretendan operar "
      "como moneda alternativa al dólar, las operaciones sin transferencia "
      "efectiva de valor, los actos del Art. 14 bis de la LOFPD y los "
      "documentos clasificados como reservados o secretos que afecten a la "
      "seguridad del Estado. Adicionalmente, se incorpora la facultad de "
      "restricción del regulador y un régimen reforzado de tipificación que "
      "elimina la defensa de neutralidad algorítmica para el nodo validador "
      "que valida transacciones con conocimiento de su ilicitud.")

    h2(doc, "VII.7 Reglamento Técnico de la Tercera Etapa")
    p(doc,
      "El reglamento técnico desarrolla la operación de la Bolsa Soberana de "
      "Productos, las categorías iniciales de productos tokenizables "
      "(incluyendo cacao CCN-51 y nacional fino de aroma, banano por "
      "variedad, camarón con trazabilidad de origen, café arábica y robusta, "
      "oro Au999 y Au995, petróleo crudo Oriente y Napo, energía "
      "hidroeléctrica y créditos de carbono certificados), el régimen del "
      "Fideicomiso Financiero Digital, las SECAS, el régimen simplificado "
      "MIPYMEs, las características técnicas del seguro de ciberseguridad y "
      "el catálogo dinámico de activos excluidos. Adicionalmente articula los "
      "regímenes del Emisor de Pago Electrónico y del Valor Negociable "
      "Endosable Digital de la Segunda Etapa.")

    # ====== VIII. SANDBOX ======
    h1(doc, "VIII. El sandbox regulatorio como puente operativo")
    p(doc,
      "El sandbox regulatorio, modelado sobre la Guía de Referencia del "
      "Sandbox de la Cámara Paraguaya de Blockchain V1.0, es el puente "
      "operativo entre la habilitación legal de la infraestructura y su "
      "operación masiva. Permite a la Junta y a la Superintendencia "
      "flexibilizar temporalmente exigencias normativas bajo supervisión "
      "intensificada y trazabilidad probatoria, generando la evidencia "
      "empírica que alimenta la normativa secundaria definitiva. El sandbox "
      "se incorpora desde el Título V de la LOFPD y se replica para el "
      "perímetro del mercado de valores en la Segunda y Tercera Etapas.")

    # ====== IX. COORDINACION ======
    h1(doc, "IX. Coordinación interinstitucional y articulación con las comisiones especializadas")
    p(doc,
      "La implementación exige coordinación sostenida entre los órganos del "
      "Estado, los reguladores sectoriales y la academia. Las principales "
      "instancias de articulación son:")
    bul(doc, "Junta de Política y Regulación Financiera y Monetaria — órgano rector normativo.")
    bul(doc, "Banco Central del Ecuador — autoridad monetaria y agente fiscal y financiero del Estado.")
    bul(doc, "Superintendencia de Compañías, Valores y Seguros — autoridad de supervisión del mercado tokenizado.")
    bul(doc, "Superintendencia de Bancos — supervisión prudencial del sistema financiero nacional.")
    bul(doc, "Corporación del Seguro de Depósitos — administradora del Fondo de Garantía del Inversionista Digital.")
    bul(doc, "Autoridad de Servicios de Confianza — acreditación de prestadores cualificados y dictamen sobre excepciones a la equivalencia funcional.")
    bul(doc, "Unidad de Análisis Financiero y Económico — recepción de reportes automatizados.")
    bul(doc, "Ministerio de Economía y Finanzas — ente rector de las finanzas públicas.")
    bul(doc, "Ministerio de Telecomunicaciones y de la Sociedad de la Información — política digital nacional.")
    bul(doc, "Procuraduría General del Estado — patrocinio estatal.")
    p(doc,
      "La estrategia legislativa se articula en paralelo con la Comisión "
      "Especializada Permanente de Soberanía, Integración y Seguridad "
      "Integral, en materia de prevención del lavado de activos y "
      "financiamiento de delitos. La sincronización con esa Comisión asegura "
      "que el régimen sancionatorio de las tres etapas guarde coherencia con "
      "el Plan Nacional de Acción Estratégico contra el lavado de activos "
      "2026-2030 y con el principio de regulación basada en riesgo, sin "
      "sustitución de la regulación por la etiqueta tecnológica del activo. "
      "Esta articulación es de buena técnica legislativa y constituye una "
      "constante de la estrategia.")

    # ====== X. HOJA DE RUTA ======
    h1(doc, "X. Hoja de ruta — cronograma generacional")
    p(doc,
      "La implementación se desarrolla en cuatro horizontes temporales que "
      "articulan los hitos legislativos con los hitos operativos:")
    h3(doc, "Años 1-2 (2026-2027) — habilitación normativa")
    p(doc,
      "Aprobación de la Ley Orgánica de la Fe Pública Digital y su Reglamento "
      "Técnico. Aprobación de la Ley Orgánica Reformatoria al COMYF, LMV y "
      "COPLAFIP (versión 2) y su Reglamento. Aprobación del régimen de "
      "Comercialización Tokenizada de la Tercera Etapa (Opción A o Opción B) "
      "con su respectivo Reglamento. Suscripción del convenio "
      "interinstitucional Banco Central del Ecuador — Ministerio de "
      "Telecomunicaciones — Corporación Nacional de Telecomunicaciones EP — "
      "Fundación Blockchain Soberana del Ecuador — Academia.")
    h3(doc, "Años 3-5 (2028-2030) — migración registral")
    p(doc,
      "Migración masiva de registros públicos al formato electrónico con "
      "equivalencia funcional plena: Notarías, Registros de la Propiedad, "
      "Registros Mercantiles, Catastros, Registro Civil, IESS. Inicio "
      "operativo del Sandbox Regulatorio con al menos tres participantes "
      "admitidos en el primer ciclo.")
    h3(doc, "Años 6-15 (2031-2040) — tokenización progresiva")
    p(doc,
      "Tokenización de Activos del Mundo Real: cacao, banano, camarón, café, "
      "flores, oro, plata, cobre, petróleo, gas, energía hidroeléctrica, "
      "créditos de carbono, derechos sobre cosechas, cuentas por cobrar y "
      "derechos fiduciarios. Operación plena de la Bolsa Soberana de "
      "Productos. Despliegue del régimen simplificado para MIPYMEs.")
    h3(doc, "Años 15-25 (2041-2051) — integración regional")
    p(doc,
      "Mercado de valores digital ecuatoriano plenamente operativo, integrado "
      "con la red soberana del Paraguay y con otros mercados regionales "
      "mediante acuerdos de reconocimiento mutuo y liquidación cross-chain en "
      "T+0 bajo principio de entrega contra pago. Consolidación generacional: "
      "el ciudadano ecuatoriano nacido en 2026 recibe toda su vida documental "
      "(partida de nacimiento, cédula, título universitario, escritura de "
      "vivienda, certificados de salud, declaraciones tributarias, herencia) "
      "en un único expediente electrónico con equivalencia funcional plena.")

    # ====== XI. INDICADORES ======
    h1(doc, "XI. Indicadores de éxito del programa")
    bul(doc, "Aprobación de la LOFPD por la Asamblea Nacional con mayoría calificada (Art. 133 CRE).")
    bul(doc, "Aprobación de la Ley Reformatoria de la Segunda Etapa con la integración de las enmiendas v2.")
    bul(doc, "Aprobación del régimen de la Tercera Etapa con las cinco enmiendas estructurales y el régimen de exclusiones.")
    bul(doc, "Dictamen favorable o no objeción de la Corte Constitucional bajo Art. 438, num. 3 CRE.")
    bul(doc, "Suscripción del convenio interinstitucional sin cargo al Presupuesto General del Estado.")
    bul(doc, "Inicio operativo del Sandbox Regulatorio con al menos tres participantes en el primer ciclo.")
    bul(doc, "Cero cuestionamientos de constitucionalidad durante los primeros 24 meses de vigencia.")
    bul(doc, "Primera tokenización de un Activo del Mundo Real ecuatoriano registrada en la Infraestructura con plena equivalencia funcional.")
    bul(doc, "Primera operación cerrada en la Bolsa Soberana de Productos.")
    bul(doc, "Primera emisión bajo régimen simplificado MIPYMEs colocada exitosamente en el mercado primario digital.")

    # ====== XII. RIESGOS Y MITIGANTES ======
    h1(doc, "XII. Riesgos y mitigantes de la implementación")
    h3(doc, "XII.1 Riesgo probatorio")
    p(doc,
      "La habilitación del mercado tokenizado sin sanción previa o paralela "
      "de la LOFPD privaría al token de fundamento legal expreso de "
      "equivalencia funcional plena. Mitigante: aprobación sincronizada o, en "
      "su defecto, incorporación de la norma puente prevista en la "
      "Disposición General Segunda de la Segunda Etapa.")
    h3(doc, "XII.2 Riesgo de captura regulatoria")
    p(doc,
      "Mitigante: gobernanza tripartita 33-33-33 con prohibición expresa de "
      "concentración mayor a un tercio de los nodos validadores por un solo "
      "sector, y verificación periódica por la Superintendencia.")
    h3(doc, "XII.3 Riesgo de invasión competencial al Banco Central")
    p(doc,
      "Mitigante: prohibición taxativa de emisión de moneda alternativa al "
      "dólar (Disposiciones Generales Cuarta y Sexta de la Segunda Etapa) y "
      "operación del Emisor de Pago Electrónico exclusivamente sobre saldos en "
      "dólares con cuenta de respaldo en el Banco Central.")
    h3(doc, "XII.4 Riesgo de afectación a la protección de datos")
    p(doc,
      "Mitigante: bloque transversal de protección de datos personales con "
      "off-chaining, twin-chain, crypto-shredding y prohibición de anotación "
      "de datos sensibles.")
    h3(doc, "XII.5 Riesgo de ciberataques al custodio digital")
    p(doc,
      "Mitigante: seguro obligatorio de ciberseguridad por al menos el diez "
      "por ciento del valor de las posiciones custodiadas, auditoría externa "
      "permanente y plan de continuidad operativa.")
    h3(doc, "XII.6 Riesgo de exposición indebida del inversionista minorista")
    p(doc,
      "Mitigante: cuota máxima de exposición por inversionista minorista, "
      "Fondo de Garantía del Inversionista Digital administrado por la "
      "Corporación del Seguro de Depósitos y mediación previa obligatoria "
      "ante la Autoridad de Servicios de Confianza.")

    # ====== XIII. RECOMENDACIONES OPERATIVAS ======
    h1(doc, "XIII. Recomendaciones operativas para la implementación")
    num(doc, "Sincronizar en una misma agenda legislativa los tres proyectos: LOFPD, Ley Reformatoria de la Segunda Etapa (v2) y régimen de la Tercera Etapa (v2 — Opción A o B).")
    num(doc, "Incluir en la Ley Reformatoria de la Segunda Etapa la Disposición General Segunda de remisión supletoria a la LOFPD.")
    num(doc, "Convocar a la Corte Constitucional para consulta facultativa (Art. 438, num. 3 CRE) sobre la compatibilidad del programa con el bloque de constitucionalidad.")
    num(doc, "Suscribir el convenio interinstitucional Banco Central — Ministerio de Telecomunicaciones — Corporación Nacional de Telecomunicaciones EP — Fundación Blockchain Soberana — Academia, con cero cargo al Presupuesto General del Estado.")
    num(doc, "No legislar el proveedor sino el principio — neutralidad tecnológica. La referencia al programa EcuaLedger Soberana debe ser descriptiva, no exclusiva.")
    num(doc, "Articular la agenda con la Comisión Especializada Permanente de Soberanía, Integración y Seguridad Integral en materia antilavado, garantizando coherencia del régimen sancionatorio en las tres etapas.")
    num(doc, "Adoptar el principio de gradualidad en la implementación operativa, con habilitación del Sandbox como puente probatorio.")
    num(doc, "Publicar el catálogo dinámico de activos no tokenizables y de actos excluidos en formato legible por máquina.")
    num(doc, "Coordinar con la Autoridad de Protección de Datos la verificación periódica del bloque transversal de protección de datos personales.")
    num(doc, "Mantener interlocución técnica permanente con el referente comparado paraguayo a través de la Fundación Blockchain Soberana del Ecuador.")

    # ====== XIV. NOTA FINAL ======
    h1(doc, "XIV. Nota final")
    p(doc,
      "La presente guía se elabora bajo la responsabilidad de la asesoría "
      "constitucional del programa EcuaLedger Soberana. Su finalidad es "
      "orientativa y contributiva, no vinculante. Las recomendaciones aquí "
      "consignadas se ofrecen a la Comisión Especializada Permanente de "
      "Régimen Económico y Tributario, a las comisiones especializadas con las "
      "que se articulan los regímenes complementarios, a las autoridades de "
      "regulación y control del sistema financiero y del mercado de valores, "
      "y a las contrapartes técnico-académicas del programa, con el propósito "
      "de facilitar el debate parlamentario y la coordinación interinstitucional.")
    p(doc,
      "El éxito del programa depende de la integridad de la secuencia "
      "legislativa, de la gobernanza tripartita de la infraestructura y de la "
      "voluntad sostenida del Estado ecuatoriano de inscribir el valor en "
      "territorio soberano con la velocidad de la red y la solidez de la "
      "Constitución.",
      italic=True)

    out_docx = DEST / "INFORME_REFORMA_COMYF_IBPP_2026_v4_LOFPD_Guia_Implementacion.docx"
    out_pdf = DEST / "INFORME_REFORMA_COMYF_IBPP_2026_v4_LOFPD_Guia_Implementacion.pdf"
    doc.save(out_docx)
    print(f"OK DOCX: {out_docx.name}")
    d2p_convert(str(out_docx), str(out_pdf))
    print(f"OK PDF : {out_pdf.name}")
    return out_docx, out_pdf


if __name__ == "__main__":
    build()
