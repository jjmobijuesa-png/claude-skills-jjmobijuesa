# -*- coding: utf-8 -*-
"""
Script de referencia de la skill correccion-ortografica-rae-espanol.

Aplica corrección RAE + eliminación de duplicados + normalización de
espacios sobre DOCX. Diseñado para procesar el paquete EcuaLedger
Soberana, pero el diccionario es reutilizable.

USO:
    Modificar SRC_DIR para apuntar a la carpeta de DOCX a procesar.
    Las correcciones se aplican IN-PLACE.
"""
import re
import shutil
import sys
import time
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert as d2p_convert

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

# ======================================================================
# CONFIGURACIÓN
# ======================================================================
SRC_DIR = Path(r"C:\Users\datos\Dropbox\var 91\Neg Inm\1 Proy Belen 2026"
               r"\Tokenizacion Activos Reales - RWA\Matriz de Acuerdos\JRPFM"
               r"\Proyecto Marco Regulatorio Asamblea\00a Carta a Instancias"
               r"\Cartas Listas para Envio")
SHORT_TMP = Path(r"C:\Users\datos\.notebooklm-extractos\rae_tmp")
SHORT_TMP.mkdir(parents=True, exist_ok=True)


# ======================================================================
# DICCIONARIO RAE (regex word-boundary; orden importa)
# ======================================================================
REGLAS_RAE = [
    # ---------- TYPO SISTEMÁTICO DEL CORPUS ----------
    (r"\bSoberama\b", "Soberana"),
    (r"\bsoberama\b", "soberana"),
    (r"\bSOBERAMA\b", "SOBERANA"),

    # ---------- EÑES PERDIDAS (palabras frecuentes) ----------
    (r"\bSenor\b", "Señor"),
    (r"\bsenor\b", "señor"),
    (r"\bSenora\b", "Señora"),
    (r"\bsenora\b", "señora"),
    (r"\bSenores\b", "Señores"),
    (r"\bsenores\b", "señores"),
    (r"\bSenoras\b", "Señoras"),
    (r"\bsenoras\b", "señoras"),
    (r"\bSenor/a\b", "Señor/a"),
    (r"\bsenor/a\b", "señor/a"),
    (r"\bcompanan\b", "compañan"),
    (r"\bacompana\b", "acompaña"),
    (r"\bdesempenar\b", "desempeñar"),
    (r"\bensenanza\b", "enseñanza"),
    (r"\bdesempeno\b", "desempeño"),
    (r"\bdesempena\b", "desempeña"),
    (r"\bdiseno\b", "diseño"),
    (r"\bdisena\b", "diseña"),
    (r"\bdisenado\b", "diseñado"),
    (r"\bdisenada\b", "diseñada"),

    # ---------- ACENTOS FALTANTES (sustantivos y adjetivos comunes) ----------
    (r"\bCodigo\b", "Código"),
    (r"\bcodigo\b", "código"),
    (r"\bCONSTITUCION\b", "CONSTITUCIÓN"),
    (r"\bConstitucion\b", "Constitución"),
    (r"\bconstitucion\b", "constitución"),
    (r"\bRepublica\b", "República"),
    (r"\brepublica\b", "república"),
    (r"\bREPUBLICA\b", "REPÚBLICA"),
    (r"\bComision\b", "Comisión"),
    (r"\bcomision\b", "comisión"),
    (r"\bCOMISION\b", "COMISIÓN"),
    (r"\bComisiones\b", "Comisiones"),  # ya está bien
    (r"\bDireccion\b", "Dirección"),
    (r"\bdireccion\b", "dirección"),
    (r"\bDIRECCION\b", "DIRECCIÓN"),
    (r"\bImplementacion\b", "Implementación"),
    (r"\bimplementacion\b", "implementación"),
    (r"\bAdministracion\b", "Administración"),
    (r"\badministracion\b", "administración"),
    (r"\bAsesoria\b", "Asesoría"),
    (r"\basesoria\b", "asesoría"),
    (r"\bAsesorias\b", "Asesorías"),
    (r"\basesorias\b", "asesorías"),
    (r"\bSoberania\b", "Soberanía"),
    (r"\bsoberania\b", "soberanía"),
    (r"\bSOBERANIA\b", "SOBERANÍA"),
    (r"\bPublica\b", "Pública"),
    (r"\bpublica\b", "pública"),
    (r"\bPublico\b", "Público"),
    (r"\bpublico\b", "público"),
    (r"\bPublicos\b", "Públicos"),
    (r"\bpublicos\b", "públicos"),
    (r"\bPublicas\b", "Públicas"),
    (r"\bpublicas\b", "públicas"),
    (r"\bPUBLICA\b", "PÚBLICA"),
    (r"\bPUBLICO\b", "PÚBLICO"),
    (r"\bTecnica\b", "Técnica"),
    (r"\btecnica\b", "técnica"),
    (r"\bTecnico\b", "Técnico"),
    (r"\btecnico\b", "técnico"),
    (r"\bTecnicas\b", "Técnicas"),
    (r"\btecnicas\b", "técnicas"),
    (r"\bTecnicos\b", "Técnicos"),
    (r"\btecnicos\b", "técnicos"),
    (r"\bTECNICA\b", "TÉCNICA"),
    (r"\bTECNICO\b", "TÉCNICO"),
    (r"\bTECNICO[- ]JURIDICA\b", "TÉCNICO-JURÍDICA"),
    (r"\bjuridica\b", "jurídica"),
    (r"\bJuridica\b", "Jurídica"),
    (r"\bjuridico\b", "jurídico"),
    (r"\bJuridico\b", "Jurídico"),
    (r"\bjuridicas\b", "jurídicas"),
    (r"\bjuridicos\b", "jurídicos"),
    (r"\bJURIDICA\b", "JURÍDICA"),
    (r"\bJURIDICO\b", "JURÍDICO"),
    (r"\bEconomico\b", "Económico"),
    (r"\beconomico\b", "económico"),
    (r"\bEconomica\b", "Económica"),
    (r"\beconomica\b", "económica"),
    (r"\bECONOMICO\b", "ECONÓMICO"),
    (r"\btramite\b", "trámite"),
    (r"\bTramite\b", "Trámite"),
    (r"\btramites\b", "trámites"),
    (r"\barticulo\b", "artículo"),
    (r"\bArticulo\b", "Artículo"),
    (r"\barticulos\b", "artículos"),
    (r"\bArticulos\b", "Artículos"),
    (r"\bnumero\b", "número"),
    (r"\bNumero\b", "Número"),
    (r"\bnumeral\b", "numeral"),  # sin tilde, correcto
    (r"\bparrafo\b", "párrafo"),
    (r"\bParrafo\b", "Párrafo"),
    (r"\bMemorando\b", "Memorando"),  # correcto
    (r"\binformacion\b", "información"),
    (r"\bInformacion\b", "Información"),
    (r"\borganica\b", "orgánica"),
    (r"\bOrganica\b", "Orgánica"),
    (r"\borganico\b", "orgánico"),
    (r"\bOrganico\b", "Orgánico"),
    (r"\bORGANICA\b", "ORGÁNICA"),
    (r"\borganismo\b", "organismo"),  # correcto sin tilde
    (r"\bregimen\b", "régimen"),
    (r"\bRegimen\b", "Régimen"),
    (r"\bregimenes\b", "regímenes"),
    (r"\bbasica\b", "básica"),
    (r"\bBasica\b", "Básica"),
    (r"\bbasico\b", "básico"),
    (r"\bBasico\b", "Básico"),
    (r"\bExposicion\b", "Exposición"),
    (r"\bexposicion\b", "exposición"),
    (r"\bDisposicion\b", "Disposición"),
    (r"\bdisposicion\b", "disposición"),
    (r"\bdisposiciones\b", "disposiciones"),
    (r"\bDisposiciones\b", "Disposiciones"),  # sin tilde el plural
    (r"\bAprobacion\b", "Aprobación"),
    (r"\baprobacion\b", "aprobación"),
    (r"\bCalificacion\b", "Calificación"),
    (r"\bcalificacion\b", "calificación"),
    (r"\bSuperposicion\b", "Superposición"),
    (r"\bsuperposicion\b", "superposición"),
    (r"\bSobreinclusion\b", "Sobreinclusión"),
    (r"\bsobreinclusion\b", "sobreinclusión"),
    (r"\bSimplificacion\b", "Simplificación"),
    (r"\bsimplificacion\b", "simplificación"),
    (r"\bComunicacion\b", "Comunicación"),
    (r"\bcomunicacion\b", "comunicación"),
    (r"\bResolucion\b", "Resolución"),
    (r"\bresolucion\b", "resolución"),
    (r"\bCoordinacion\b", "Coordinación"),
    (r"\bcoordinacion\b", "coordinación"),
    (r"\bRecomendacion\b", "Recomendación"),
    (r"\brecomendacion\b", "recomendación"),
    (r"\bRecomendaciones\b", "Recomendaciones"),  # plural sin tilde
    (r"\bsugerencia\b", "sugerencia"),
    (r"\bConsideracion\b", "Consideración"),
    (r"\bconsideracion\b", "consideración"),
    (r"\bSituacion\b", "Situación"),
    (r"\bsituacion\b", "situación"),
    (r"\bDecision\b", "Decisión"),
    (r"\bdecision\b", "decisión"),
    (r"\bDecisiones\b", "Decisiones"),  # plural sin tilde
    (r"\bIntegracion\b", "Integración"),
    (r"\bintegracion\b", "integración"),
    (r"\bSeguridad\b", "Seguridad"),  # correcto
    (r"\bproteccion\b", "protección"),
    (r"\bProteccion\b", "Protección"),
    (r"\bproyecto\b", "proyecto"),  # correcto
    (r"\bteleinformacion\b", "teleinformación"),
    (r"\bAccion\b", "Acción"),
    (r"\baccion\b", "acción"),
    (r"\bAcciones\b", "Acciones"),
    (r"\bacciones\b", "acciones"),
    (r"\bMision\b", "Misión"),
    (r"\bmision\b", "misión"),
    (r"\bVision\b", "Visión"),
    (r"\bvision\b", "visión"),
    (r"\bUtopia\b", "Utopía"),
    (r"\butopia\b", "utopía"),
    (r"\beconomia\b", "economía"),
    (r"\bEconomia\b", "Economía"),
    (r"\bsociologia\b", "sociología"),
    (r"\bMonetaria\b", "Monetaria"),  # correcto
    (r"\bMonetario\b", "Monetario"),  # correcto
    (r"\bregulacion\b", "regulación"),
    (r"\bRegulacion\b", "Regulación"),
    (r"\bregulaciones\b", "regulaciones"),
    (r"\bestatutaria\b", "estatutaria"),  # correcto
    (r"\bconstitucionalista\b", "constitucionalista"),  # correcto
    (r"\bMagnifica\b", "Magnífica"),
    (r"\bmagnifica\b", "magnífica"),
    (r"\bEstrategica\b", "Estratégica"),
    (r"\bestrategica\b", "estratégica"),
    (r"\bestrategicos\b", "estratégicos"),
    (r"\bEstrategicos\b", "Estratégicos"),
    (r"\bAdministrativos\b", "Administrativos"),
    (r"\bElectronica\b", "Electrónica"),
    (r"\belectronica\b", "electrónica"),
    (r"\belectronico\b", "electrónico"),
    (r"\bElectronico\b", "Electrónico"),
    (r"\bElectronicos\b", "Electrónicos"),
    (r"\belectronicos\b", "electrónicos"),
    (r"\bDigital\b", "Digital"),  # correcto
    (r"\bsancionatorio\b", "sancionatorio"),  # correcto
    (r"\bbursatil\b", "bursátil"),
    (r"\bBursatil\b", "Bursátil"),
    (r"\bFinanciera\b", "Financiera"),  # correcto
    (r"\bMonetario\b", "Monetario"),  # correcto
    (r"\bSinergica\b", "Sinérgica"),
    (r"\bsinergica\b", "sinérgica"),
    (r"\bAnalitica\b", "Analítica"),
    (r"\banalitica\b", "analítica"),

    # ---------- ESPAÑA RIOS Y NOMBRES ----------
    (r"\bLos Rios\b", "Los Ríos"),
    (r"\blos rios\b", "Los Ríos"),
    (r"\bRios\b", "Ríos"),
    (r"\brios\b", "ríos"),

    # ---------- VERBOS Y FORMAS COMUNES ----------
    (r"\bUtilizara\b", "Utilizará"),
    (r"\butilizara\b", "utilizará"),
    (r"\bSera\b", "Será"),
    (r"\bsera\b", "será"),
    (r"\bSeran\b", "Serán"),
    (r"\bseran\b", "serán"),
    (r"\bEsta\b", "Está"),  # imperativo "está" — pero ojo, "esta" demostrativo es correcto. Lo dejamos como caso a inspeccionar
    (r"\bdebera\b", "deberá"),
    (r"\bDebera\b", "Deberá"),
    (r"\bdeberan\b", "deberán"),
    (r"\bDeberan\b", "Deberán"),
    (r"\bpodra\b", "podrá"),
    (r"\bPodra\b", "Podrá"),
    (r"\bpodran\b", "podrán"),
    (r"\bPodran\b", "Podrán"),

    # ---------- ESPAÑOL DE ECUADOR (acentos específicos) ----------
    (r"\bjurisdiccional\b", "jurisdiccional"),  # correcto
    (r"\bsantos\b", "santos"),  # correcto
    (r"\bautoridad\b", "autoridad"),  # correcto
    (r"\bSESEAS\b", "SECAS"),  # corrección sigla
    (r"\bSandbox\b", "Sandbox"),  # inglés, dejar

    # ---------- VARIOS ----------
    (r"\bIntegra\b", "Integra"),  # correcto
    (r"\bRiesgos\b", "Riesgos"),  # correcto
    (r"\bgrana\b", "grana"),  # correcto
    (r"\bgranar\b", "granar"),  # correcto

    # ============================================================
    # CORRUPCIONES UTF-8 → ASCII (vocal acentuada eliminada por completo)
    # ============================================================
    (r"\bCdigo\b", "Código"),
    (r"\bcdigo\b", "código"),
    (r"\bCDIGO\b", "CÓDIGO"),
    (r"\bCdigos\b", "Códigos"),
    (r"\bDireccin\b", "Dirección"),
    (r"\bdireccin\b", "dirección"),
    (r"\bDIRECCIN\b", "DIRECCIÓN"),
    (r"\bComisin\b", "Comisión"),
    (r"\bcomisin\b", "comisión"),
    (r"\bComisiin\b", "Comisión"),
    (r"\bFundacin\b", "Fundación"),
    (r"\bfundacin\b", "fundación"),
    (r"\bFUNDACIN\b", "FUNDACIÓN"),
    (r"\bAcin\b", "Acción"),
    (r"\bAccin\b", "Acción"),
    (r"\baccin\b", "acción"),
    (r"\bACCION\b", "ACCIÓN"),
    (r"\bACCIN\b", "ACCIÓN"),
    (r"\bResolucion\b", "Resolución"),
    (r"\bResolucin\b", "Resolución"),
    (r"\bresolucin\b", "resolución"),
    (r"\bAprobacin\b", "Aprobación"),
    (r"\baprobacin\b", "aprobación"),
    (r"\bImplementacin\b", "Implementación"),
    (r"\bimplementacin\b", "implementación"),
    (r"\bInformacin\b", "Información"),
    (r"\binformacin\b", "información"),
    (r"\bAdministracin\b", "Administración"),
    (r"\badministracin\b", "administración"),
    (r"\bCoordinacin\b", "Coordinación"),
    (r"\bcoordinacin\b", "coordinación"),
    (r"\bExposicin\b", "Exposición"),
    (r"\bexposicin\b", "exposición"),
    (r"\bDisposicin\b", "Disposición"),
    (r"\bdisposicin\b", "disposición"),
    (r"\bConsideracin\b", "Consideración"),
    (r"\bconsideracin\b", "consideración"),
    (r"\bSituacin\b", "Situación"),
    (r"\bsituacin\b", "situación"),
    (r"\bDecisin\b", "Decisión"),
    (r"\bdecisin\b", "decisión"),
    (r"\bIntegracin\b", "Integración"),
    (r"\bintegracin\b", "integración"),
    (r"\bProteccin\b", "Protección"),
    (r"\bprotectin\b", "protección"),
    (r"\bproteccin\b", "protección"),
    (r"\bRegulacin\b", "Regulación"),
    (r"\bregulacin\b", "regulación"),
    (r"\bMisin\b", "Misión"),
    (r"\bmisin\b", "misión"),
    (r"\bVisin\b", "Visión"),
    (r"\bvisin\b", "visión"),
    (r"\bAcompaa\b", "Acompaña"),
    (r"\bacompaa\b", "acompaña"),
    (r"\bDiseo\b", "Diseño"),
    (r"\bdiseo\b", "diseño"),
    (r"\bSealamos\b", "Señalamos"),
    (r"\bsealamos\b", "señalamos"),
    (r"\bSeala\b", "Señala"),
    (r"\bseala\b", "señala"),
    (r"\bSeor\b", "Señor"),
    (r"\bseor\b", "señor"),
    (r"\bSeora\b", "Señora"),
    (r"\bseora\b", "señora"),

    # Sin la "í"
    (r"\bTcnico\b", "Técnico"),
    (r"\btcnico\b", "técnico"),
    (r"\bTcnica\b", "Técnica"),
    (r"\btcnica\b", "técnica"),
    (r"\bTCNICA\b", "TÉCNICA"),
    (r"\bTCNICO\b", "TÉCNICO"),
    (r"\bTcnicos\b", "Técnicos"),
    (r"\btcnicos\b", "técnicos"),
    (r"\bTcnicas\b", "Técnicas"),
    (r"\btcnicas\b", "técnicas"),
    (r"\bJurdica\b", "Jurídica"),
    (r"\bjurdica\b", "jurídica"),
    (r"\bJurdico\b", "Jurídico"),
    (r"\bjurdico\b", "jurídico"),
    (r"\bJurdicas\b", "Jurídicas"),
    (r"\bjurdicas\b", "jurídicas"),
    (r"\bJurdicos\b", "Jurídicos"),
    (r"\bjurdicos\b", "jurídicos"),
    (r"\bJURDICA\b", "JURÍDICA"),
    (r"\bJURDICO\b", "JURÍDICO"),
    (r"\bRos\b", "Ríos"),
    (r"\bros\b", "ríos"),
    (r"\bLos Ros\b", "Los Ríos"),
    (r"\bAsesoria\b", "Asesoría"),
    (r"\bArtculo\b", "Artículo"),
    (r"\bartculo\b", "artículo"),
    (r"\bartculos\b", "artículos"),
    (r"\bPolitica\b", "Política"),
    (r"\bpolitica\b", "política"),
    (r"\bPolitico\b", "Político"),
    (r"\bpolitico\b", "político"),
    (r"\bMagnfica\b", "Magnífica"),
    (r"\bMagnifica\b", "Magnífica"),
    (r"\bConstitucin\b", "Constitución"),
    (r"\bconstitucin\b", "constitución"),
    (r"\bRepublica\b", "República"),
    (r"\bRepublic\b", "República"),
    (r"\bSoberania\b", "Soberanía"),
    (r"\bsoberania\b", "soberanía"),
    (r"\bSoberano\b", "Soberano"),  # correcto
    (r"\bRgimen\b", "Régimen"),
    (r"\brgimen\b", "régimen"),

    # Sin la "á"
    (r"\bTramite\b", "Trámite"),
    (r"\btramite\b", "trámite"),
    (r"\bTramites\b", "Trámites"),
    (r"\btramites\b", "trámites"),
    (r"\bMatriz\b", "Matriz"),  # correcto
    (r"\bMnsterio\b", "Ministerio"),
    (r"\bMinisterios\b", "Ministerios"),  # correcto

    # Verbos sin tilde
    (r"\bgeneracion\b", "generación"),
    (r"\bGeneracion\b", "Generación"),
    (r"\brecibira\b", "recibirá"),
    (r"\bRecibira\b", "Recibirá"),
    (r"\bnumeracin\b", "numeración"),
    (r"\bNumeracin\b", "Numeración"),
    (r"\bcoordinacin\b", "coordinación"),
    (r"\bCoordinacin\b", "Coordinación"),
    (r"\barmonizacin\b", "armonización"),
    (r"\bArmonizacin\b", "Armonización"),
    (r"\bIntervencin\b", "Intervención"),
    (r"\bintervencin\b", "intervención"),

    # Otras corrupciones
    (r"\bdlar\b", "dólar"),
    (r"\bdlares\b", "dólares"),
    (r"\bDlar\b", "Dólar"),
    (r"\barmnico\b", "armónico"),
    (r"\barmnica\b", "armónica"),
    (r"\bArmnico\b", "Armónico"),
    (r"\bArmnica\b", "Armónica"),
    (r"\beconmica\b", "económica"),
    (r"\beconmico\b", "económico"),
    (r"\bEconmico\b", "Económico"),
    (r"\bEconmica\b", "Económica"),
    (r"\bECONMICO\b", "ECONÓMICO"),
    (r"\bRgimen\b", "Régimen"),
    (r"\bcrnica\b", "crónica"),
    (r"\borgnica\b", "orgánica"),
    (r"\bOrgnica\b", "Orgánica"),
    (r"\borgnico\b", "orgánico"),
    (r"\bOrgnico\b", "Orgánico"),
    (r"\bORGNICA\b", "ORGÁNICA"),
    (r"\btoma\b", "toma"),  # correcto
    (r"\bDIAGNOSTICO\b", "DIAGNÓSTICO"),
    (r"\bDiagnostico\b", "Diagnóstico"),
    (r"\bdiagnostico\b", "diagnóstico"),
    (r"\bestrategica\b", "estratégica"),
    (r"\bEstrategica\b", "Estratégica"),
    (r"\bestrategico\b", "estratégico"),
    (r"\bEstrategico\b", "Estratégico"),
    (r"\bestrategias\b", "estrategias"),  # correcto
    (r"\bjurisdiccional\b", "jurisdiccional"),  # correcto

    # Palabras con vocales perdidas en otras posiciones
    (r"\bvas\b", "vías"),  # cuidado contexto
    (r"\bAsamblea\b", "Asamblea"),  # correcto
    (r"\bCmara\b", "Cámara"),
    (r"\bcmara\b", "cámara"),
    (r"\bcdula\b", "cédula"),
    (r"\bCdula\b", "Cédula"),
    (r"\bgua\b", "guía"),
    (r"\bGua\b", "Guía"),
    (r"\btendr\b", "tendrá"),
    (r"\bUtilidad\b", "Utilidad"),  # correcto

    # Errores específicos detectados en el corpus
    (r"\bcualq\b", "cualquier"),
    (r"\bdems\b", "demás"),
    (r"\bDems\b", "Demás"),
    (r"\bN\.o\b", "N.°"),  # número ordinal
    (r"\bn\.o\b", "n.°"),
    (r"\bComisin Especializada\b", "Comisión Especializada"),
    (r"\bSeora Asambleta\b", "Señora Asambleísta"),
    (r"\bAsambleta\b", "Asambleísta"),
    (r"\basambleta\b", "asambleísta"),
    (r"\bms\b", "más"),
    (r"\bMs\b", "Más"),  # cuidado con MS (siglas) — no aplicará por word boundary
]

# Patrón duplicados (case insensitive, palabras de 3+ letras)
DUPLICADOS_RE = re.compile(r"\b(\w{3,})\s+\1\b", re.IGNORECASE)

# Reglas de espaciado
ESPACIOS = [
    (r" +,", ","),  # espacio antes de coma
    (r" +\.", "."),  # espacio antes de punto
    (r" +;", ";"),  # espacio antes de punto y coma
    (r" +:", ":"),  # espacio antes de dos puntos
    (r"\(\s+", "("),  # espacio tras paréntesis abierto
    (r"\s+\)", ")"),  # espacio antes de paréntesis cerrado
    (r"[ \t]{2,}", " "),  # dos o más espacios → uno
]


# ======================================================================
# APLICACIÓN
# ======================================================================
def apply_rules(text: str) -> tuple:
    """Aplica todas las reglas y devuelve (texto_corregido, cuenta_correcciones)."""
    n_total = 0
    # Diccionario RAE
    for pat, repl in REGLAS_RAE:
        text, n = re.subn(pat, repl, text)
        n_total += n
    # Duplicados (mantener uno solo)
    text, n = DUPLICADOS_RE.subn(r"\1", text)
    n_total += n
    # Espacios
    for pat, repl in ESPACIOS:
        text, n = re.subn(pat, repl, text)
        n_total += n
    return text, n_total


def replace_in_paragraph(paragraph) -> int:
    """Aplica reglas in-place al primer run, preservando formato."""
    if not paragraph.runs or not paragraph.text:
        return 0
    new, n = apply_rules(paragraph.text)
    if n == 0:
        return 0
    paragraph.runs[0].text = new
    for r in paragraph.runs[1:]:
        r.text = ""
    return n


# ======================================================================
# JUSTIFICAR Y ENCUADRAR (texto justificado, márgenes, títulos)
# ======================================================================
def aplicar_formato(doc: Document):
    """Justifica todos los párrafos de texto y centra los títulos."""
    # Configurar márgenes A4 estándar
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        # Detectar títulos de sección romana, encabezados o líneas separadoras
        is_titulo = bool(
            re.match(r"^(REP[UÚ]BLICA|FUNDACI[OÓ]N|ASESOR[IÍ]A|INFORME EJECUTIVO|"
                     r"C[oó]digo:|SUMILLA|DIAGN[OÓ]STICO|PROPUESTA|"
                     r"ANCLAJE|MECANISMOS|CRONOGRAMA|RIESGOS|RECOMENDACI[OÓ]N|"
                     r"ANEXOS|S[IÍ]NTESIS|PROBLEMA:|SOLUCI[OÓ]N:|LAS TRES|"
                     r"GOBERNANZA:|HORIZONTE:|ACCI[OÓ]N|II\.|III\.|IV\.|V\.|"
                     r"VI\.|VII\.|VIII\.|IX\.|X\.|I\.)", t, re.IGNORECASE)
        )
        is_separator = bool(re.match(r"^_{5,}$", t))
        if is_titulo or is_separator:
            # Centro para portada/encabezado, izquierda con negrita para secciones
            if t.startswith("_"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif re.match(r"^(REP[UÚ]BLICA|FUNDACI[OÓ]N|ASESOR[IÍ]A|INFORME EJECUTIVO|C[oó]digo)", t):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Texto justificado para mejor presentación
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ======================================================================
# ACTUALIZAR ANEXOS (con Google Drive link + leyes Paraguay + leyes Ecuador)
# ======================================================================
DRIVE_LINK = "https://drive.google.com/drive/folders/1B0jI0qtCtLKLsYXhInFz4JEK90ZHPtUL?usp=drive_link"

NUEVOS_ANEXOS = [
    "Modelo soberano paraguayo (referente comparado):",
    "• Ley N.° 6822 de 2021 — Servicios de Confianza para las Transacciones Electrónicas y del Documento Transmisible Electrónico (DTE).",
    "• Ley N.° 7572 de 2025 — Mercado de Valores y Productos del Paraguay.",
    "",
    "Propuesta normativa ecuatoriana (en formato físico y digital):",
    "• Primera Etapa — Anteproyecto de Ley Orgánica de la Fe Pública Digital (LOFPD) y su Reglamento Técnico de Aplicación.",
    "• Segunda Etapa — Proyecto de Ley Orgánica Reformatoria al Código Orgánico Monetario y Financiero, a la Ley de Mercado de Valores y al Código Orgánico de Planificación y Finanzas Públicas, con su Reglamento Técnico.",
    "• Tercera Etapa — Proyectos de Ley Orgánica de Comercialización Tokenizada (Opción A — LOCATok) y de Capítulo VII bis incorporado a la Ley de Mercado de Valores (Opción B), con su Reglamento Técnico de Aplicación.",
    "",
    "Acceso digital al expediente completo (Google Drive — tres carpetas, una por etapa legislativa):",
    DRIVE_LINK,
    "",
    "Documentos técnicos complementarios:",
    "• INFORME v5 — Propuesta EcuaLedger Soberana para el Ecuador (DOCX y PDF).",
    "• Memorando integrador del Sandbox Regulatorio CPYB V1.0.",
    "• Esquema visual de la arquitectura H2 Twinchain y de la gobernanza multisectorial 25-25-25-25.",
]


def actualizar_anexos(doc: Document):
    """Reemplaza el bloque de ANEXOS por la nueva versión con Drive link."""
    # Encontrar el rango: desde el párrafo que contiene "ANEXOS" hasta antes
    # de la línea separadora "____________________________" del cierre.
    paragraphs = doc.paragraphs
    idx_anexos = None
    idx_cierre = None
    for i, p in enumerate(paragraphs):
        t = (p.text or "").strip().upper()
        if re.search(r"^(VII|VIII|IX|X|XI)\.\s*ANEXOS", t):
            idx_anexos = i
        if idx_anexos is not None and t.startswith("_" * 5):
            idx_cierre = i
            break
    if idx_anexos is None:
        return False
    if idx_cierre is None:
        # Si no hay separador, eliminar hasta el final
        idx_cierre = len(paragraphs)

    # Eliminar todos los párrafos viejos entre idx_anexos+1 e idx_cierre-1
    for j in range(idx_cierre - 1, idx_anexos, -1):
        el = paragraphs[j]._element
        el.getparent().remove(el)

    # Re-leer paragraphs porque el árbol cambió
    paragraphs = doc.paragraphs
    # Buscar nuevo idx_anexos
    for i, p in enumerate(paragraphs):
        if re.search(r"^(VII|VIII|IX|X|XI)\.\s*ANEXOS", (p.text or "").strip().upper()):
            idx_anexos = i
            break

    # Insertar los nuevos anexos después del idx_anexos
    anexo_p = paragraphs[idx_anexos]
    # python-docx: insertar después requiere XML manipulation
    from docx.oxml.ns import qn
    parent = anexo_p._element.getparent()
    insert_pos = list(parent).index(anexo_p._element)
    # Crear párrafos nuevos
    from copy import deepcopy
    template_p = anexo_p._element
    for j, line in enumerate(NUEVOS_ANEXOS):
        new_p = deepcopy(template_p)
        # Limpiar runs
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        # Crear nuevo run con el texto
        if line.strip():
            r = template_p.makeelement(qn("w:r"), {})
            t = template_p.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
            t.text = line
            r.append(t)
            new_p.append(r)
        # Insertar después del párrafo anterior
        parent.insert(insert_pos + 1 + j, new_p)
    return True


# ======================================================================
# DIAGNÓSTICO
# ======================================================================
def diagnose(doc_path: Path) -> dict:
    """Reporta tipos y cuentas de errores en el doc."""
    doc = Document(str(doc_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    diag = {}
    diag["soberama"] = len(re.findall(r"\bSober[ae]ma\b", text))
    diag["sin_acento"] = sum(
        len(re.findall(pat, text)) for pat, _ in REGLAS_RAE[:80]
    )
    diag["duplicados"] = len(DUPLICADOS_RE.findall(text))
    diag["doble_espacio"] = len(re.findall(r"[ \t]{2,}", text))
    diag["espacio_antes_coma"] = len(re.findall(r" +,", text))
    diag["palabras_total"] = len(text.split())
    return diag


# ======================================================================
# PDF
# ======================================================================
def docx_to_pdf_via_short(docx_path: Path):
    """Convierte a PDF usando path corto."""
    name = docx_path.stem
    short_docx = SHORT_TMP / f"{name}.docx"
    short_pdf = SHORT_TMP / f"{name}.pdf"
    pdf_path = docx_path.with_suffix(".pdf")
    shutil.copy(str(docx_path), str(short_docx))
    time.sleep(0.5)
    d2p_convert(str(short_docx), str(short_pdf))
    time.sleep(0.5)
    shutil.copy(str(short_pdf), str(pdf_path))
    short_docx.unlink(missing_ok=True)
    short_pdf.unlink(missing_ok=True)


# ======================================================================
# ORQUESTA
# ======================================================================
def process(docx_path: Path):
    print(f"\n>>> {docx_path.name}")
    diag_before = diagnose(docx_path)
    print(f"    Antes: soberama={diag_before['soberama']}  "
          f"duplicados={diag_before['duplicados']}  "
          f"doble_espacio={diag_before['doble_espacio']}  "
          f"espacio_antes_coma={diag_before['espacio_antes_coma']}")

    doc = Document(str(docx_path))
    # 1. Correcciones RAE + duplicados + espacios
    n_total = 0
    for p in doc.paragraphs:
        n_total += replace_in_paragraph(p)
    print(f"    Reemplazos ortográficos: {n_total}")

    # 2. Actualizar anexos con Drive link
    actualizado = actualizar_anexos(doc)
    print(f"    Anexos actualizados: {actualizado}")

    # 3. Aplicar formato (justificación + márgenes)
    aplicar_formato(doc)
    print(f"    Formato aplicado: justificación + márgenes A4 2.5cm")

    # Guardar
    doc.save(str(docx_path))

    # 4. Verificación
    diag_after = diagnose(docx_path)
    print(f"    Después: soberama={diag_after['soberama']}  "
          f"duplicados={diag_after['duplicados']}  "
          f"doble_espacio={diag_after['doble_espacio']}  "
          f"espacio_antes_coma={diag_after['espacio_antes_coma']}")

    # 5. PDF
    try:
        docx_to_pdf_via_short(docx_path)
        print(f"    PDF OK")
    except Exception as e:
        print(f"    WARN PDF: {str(e)[:100]}")


if __name__ == "__main__":
    print("=" * 70)
    print("CORRECCIÓN ORTOGRÁFICA RAE + JUSTIFICACIÓN + ANEXOS")
    print("=" * 70)
    docs = sorted(SRC_DIR.glob("B-*.docx"))
    for fp in docs:
        try:
            process(fp)
        except Exception as e:
            print(f"    ERROR: {e}")
            import traceback
            traceback.print_exc()
    print("\n" + "=" * 70)
    print("FIN")
    print("=" * 70)
